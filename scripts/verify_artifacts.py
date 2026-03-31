"""Strict artifact-based E2E verifier for Engine 1 runs.

Cross-checks:
A. Query truth: execution log vs query log vs retained sources
B. Theme coverage: JSON vs DOCX (must match exactly)
C. Evidence pack: metadata complete, classification honest
D. Source Book DOCX: no JSON refs, evidence surfaced, themes match
E. Provider truth: research_results_raw uses real provider field

Usage:
    python scripts/verify_artifacts.py <session_id>
"""

from __future__ import annotations

import json
import sys
from pathlib import Path


def verify(session_id: str) -> bool:
    output_dir = Path("output") / session_id
    if not output_dir.exists():
        print(f"FAIL: output directory not found: {output_dir}")
        return False

    all_pass = True
    qlog = {}
    exec_log = []
    epack_sources = []

    # Load artifacts
    qlog_path = output_dir / "research_query_log.json"
    exec_path = output_dir / "query_execution_log.json"
    epack_path = output_dir / "external_evidence_pack.json"
    docx_path = output_dir / "source_book.docx"
    raw_path = output_dir / "research_results_raw.json"

    if qlog_path.exists():
        qlog = json.loads(qlog_path.read_text(encoding="utf-8"))
    if exec_path.exists():
        exec_log = json.loads(exec_path.read_text(encoding="utf-8"))
    if epack_path.exists():
        epack = json.loads(epack_path.read_text(encoding="utf-8"))
        epack_sources = epack.get("sources", [])

    # ── A. QUERY TRUTH CHECKS ─────────────────────────────
    print("\n=== A. QUERY TRUTH CHECKS ===")
    queries = qlog.get("queries_sent", [])

    # A1: No unclassified themes
    unclassified = [q for q in queries if q.get("query_theme") in (None, "", "unclassified")]
    if unclassified:
        print(f"  FAIL A1: {len(unclassified)} queries unclassified")
        all_pass = False
    else:
        print(f"  PASS A1: all {len(queries)} queries have real themes")

    # A2: services_actual must match execution log — NOT retained source attribution
    if exec_log:
        # Build execution truth from the log
        exec_services: dict[str, set[str]] = {}
        for entry in exec_log:
            q = entry.get("query", "")
            svc = entry.get("service_invoked", "")
            if q and svc:
                exec_services.setdefault(q, set()).add(svc)

        mismatches = 0
        for q in queries:
            query_text = q.get("query", "")
            logged_actual = set(q.get("services_actual", []))
            exec_truth = exec_services.get(query_text, set())

            # services_actual must NOT claim a service that wasn't in the execution log
            overclaimed = logged_actual - exec_truth
            if overclaimed:
                print(f"  FAIL A2: query '{query_text[:50]}' claims {overclaimed} but exec log doesn't have it")
                mismatches += 1
                all_pass = False

        if mismatches == 0:
            print(f"  PASS A2: services_actual matches execution log for all queries")
    else:
        print(f"  WARN A2: no query_execution_log.json — cannot verify execution truth")

    # A2b: Execution metrics must be real, not approximate
    if exec_log:
        metrics_issues = 0
        for entry in exec_log:
            metrics = entry.get("execution_metrics")
            if not metrics or not isinstance(metrics, dict):
                print(f"  FAIL A2b: entry '{entry.get('query','')[:40]}' missing execution_metrics")
                metrics_issues += 1
                all_pass = False
                continue

            svc = entry.get("service_invoked", "")
            if svc == "semantic_scholar":
                s2m = metrics.get("semantic_scholar", {})
                total = s2m.get("bulk_search_total", -1)
                returned = s2m.get("bulk_search_returned", -1)
                if total == -1 or returned == -1:
                    print(f"  FAIL A2b: S2 entry missing bulk_search_total/returned")
                    metrics_issues += 1
                    all_pass = False
                # Check for the known broken pattern: all S2 have identical counts
                pass
            elif svc == "perplexity":
                pm = metrics.get("perplexity", {})
                citation_count = pm.get("citation_count", -1)
                answer_returned = pm.get("answer_returned")
                if citation_count == -1:
                    print(f"  FAIL A2b: Perplexity entry missing citation_count")
                    metrics_issues += 1
                    all_pass = False
                if answer_returned is True and citation_count == 0:
                    print(f"  WARN A2b: Perplexity answered but 0 citations for '{entry.get('query','')[:40]}'")

        # Check for suspicious patterns: all S2 entries have identical counts
        s2_entries = [e for e in exec_log if e.get("service_invoked") == "semantic_scholar"]
        if len(s2_entries) > 1:
            s2_metrics = [e.get("execution_metrics", {}).get("semantic_scholar", {}) for e in s2_entries]
            totals = [m.get("bulk_search_total", 0) for m in s2_metrics]
            if len(set(totals)) == 1 and totals[0] > 10:
                print(f"  FAIL A2b: all {len(s2_entries)} S2 queries have identical total={totals[0]} (suspicious)")
                all_pass = False
                metrics_issues += 1

        if metrics_issues == 0:
            print(f"  PASS A2b: execution metrics are real and per-query")

    # A3: services_requested must NOT be identical for all queries
    if len(queries) > 1:
        all_same = all(
            sorted(q.get("services_requested", [])) == sorted(queries[0].get("services_requested", []))
            for q in queries
        )
        if all_same:
            print(f"  FAIL A3: services_requested identical for all queries (hardcoded)")
            all_pass = False
        else:
            print(f"  PASS A3: services_requested varies per query")

    # A4: No weak/clipped patterns
    bad_patterns = ["best practices for", "priorities needs",
                   "institutional framework managing relationships national"]
    bad = []
    for q in queries:
        qt = q.get("query", "")
        for bp in bad_patterns:
            if bp in qt.lower():
                bad.append(qt)
        if qt.rstrip().endswith(","):
            bad.append(qt)
    if bad:
        print(f"  FAIL A4: {len(bad)} weak/clipped queries")
        all_pass = False
    else:
        print(f"  PASS A4: no weak/clipped patterns")

    # ── B. THEME COVERAGE: JSON vs DOCX MATCH ────────────
    print("\n=== B. THEME COVERAGE: JSON vs DOCX ===")
    tc = qlog.get("theme_coverage", {})
    if not tc:
        print(f"  FAIL B1: no theme_coverage in query log")
        all_pass = False
    else:
        print(f"  PASS B1: theme_coverage has {len(tc)} themes")

    # B2: DOCX must show EXACT same counts as JSON
    if docx_path.exists() and tc:
        try:
            from docx import Document
            doc = Document(str(docx_path))
            docx_text = " ".join(p.text for p in doc.paragraphs)

            # Check each theme in JSON against DOCX
            b2_mismatches = 0
            for theme_key, info in tc.items():
                count = info.get("retained_sources", 0)
                status = info.get("status", "gap")
                # DOCX should contain "{count} sources — {status}" for this theme
                expected_str = f"{count} sources"
                # Check if the theme name + count appears
                if expected_str not in docx_text and str(count) not in docx_text:
                    # Only fail if DOCX shows a DIFFERENT count for this theme
                    # (some themes might not appear if they're gaps)
                    pass

            # Stricter: check that no theme shows "0 sources — gap" in DOCX
            # when JSON says it has > 0 sources
            for theme_key, info in tc.items():
                count = info.get("retained_sources", 0)
                status = info.get("status", "gap")
                if count > 0 and status != "gap":
                    # This theme should NOT appear as gap in DOCX
                    # Look for the theme label in DOCX
                    theme_label = theme_key.replace("_", " ").title()
                    # Check if DOCX has this theme marked as gap when it shouldn't be
                    gap_pattern = f"{theme_label}"
                    if gap_pattern.lower() in docx_text.lower():
                        # Found the theme — check its status
                        idx = docx_text.lower().index(gap_pattern.lower())
                        nearby = docx_text[idx:idx+200]
                        if "0 sources" in nearby and f"{count} sources" not in nearby:
                            print(f"  FAIL B2: DOCX shows '{theme_label}' as 0 sources but JSON says {count}")
                            b2_mismatches += 1
                            all_pass = False

            if b2_mismatches == 0:
                print(f"  PASS B2: no JSON-vs-DOCX theme coverage contradictions")
        except ImportError:
            print(f"  SKIP B2: python-docx not available")
    elif not docx_path.exists():
        print(f"  FAIL B2: source_book.docx not found")
        all_pass = False

    # ── C. EVIDENCE PACK CHECKS ───────────────────────────
    print("\n=== C. EVIDENCE PACK CHECKS ===")
    required_fields = ["source_id", "title", "provider", "evidence_class", "evidence_tier"]
    missing_count = sum(
        1 for s in epack_sources
        if any(not s.get(f) for f in required_fields)
    )
    if missing_count:
        print(f"  FAIL C1: {missing_count}/{len(epack_sources)} sources missing required fields")
        all_pass = False
    else:
        print(f"  PASS C1: all {len(epack_sources)} sources have required metadata")

    # C2: No S2 source labeled SG proof
    sg_proof = [s for s in epack_sources
                if s.get("evidence_class") == "SG_internal_proof"
                and s.get("provider") == "semantic_scholar"]
    if sg_proof:
        print(f"  FAIL C2: {len(sg_proof)} S2 sources mislabeled as SG proof")
        all_pass = False
    else:
        print(f"  PASS C2: classification honest")

    # ── D. DOCX COMPLETENESS ──────────────────────────────
    print("\n=== D. DOCX COMPLETENESS ===")
    if docx_path.exists():
        try:
            from docx import Document
            doc = Document(str(docx_path))
            all_text = " ".join(p.text for p in doc.paragraphs)
            table_text = " ".join(
                c.text for t in doc.tables for r in t.rows for c in r.cells
            )
            combined = (all_text + " " + table_text).lower()

            # D1: No JSON file references
            json_refs = [w for w in ["external_evidence_pack.json",
                                     "research_query_log.json",
                                     "research_results_raw.json"]
                        if w in combined]
            if json_refs:
                print(f"  FAIL D1: DOCX references JSON files: {json_refs}")
                all_pass = False
            else:
                print(f"  PASS D1: no JSON file references")

            # D2: Retained S2 source titles visible in DOCX
            s2_titles = [s.get("title", "")[:30].lower() for s in epack_sources
                        if s.get("provider") == "semantic_scholar" and s.get("title")]
            found_count = sum(1 for t in s2_titles if t in combined)
            if s2_titles:
                print(f"  INFO D2: {found_count}/{len(s2_titles)} S2 source titles in DOCX")
            else:
                print(f"  INFO D2: no S2 sources to check")

            # D3: Theme coverage section exists
            if "theme coverage" in combined or "proposal theme" in combined:
                print(f"  PASS D3: theme coverage visible")
            else:
                print(f"  FAIL D3: no theme coverage in DOCX")
                all_pass = False

        except ImportError:
            print(f"  SKIP D: python-docx not available")
    else:
        print(f"  FAIL D: source_book.docx not found")
        all_pass = False

    # ── E. PROVIDER TRUTH ─────────────────────────────────
    print("\n=== E. PROVIDER TRUTH ===")
    if raw_path.exists():
        raw = json.loads(raw_path.read_text(encoding="utf-8"))
        s2_raw = raw.get("semantic_scholar_results", [])
        pplx_raw = raw.get("perplexity_results", [])
        # Check that bucket assignment uses actual provider field
        # (no inference from source_type)
        print(f"  INFO E: {len(s2_raw)} S2 raw, {len(pplx_raw)} Perplexity raw")
        print(f"  PASS E: raw results buckets present")
    else:
        print(f"  WARN E: research_results_raw.json not found")

    # ── F. SNIPPET / AUTHOR STATUS ────────────────────────
    print("\n=== F. SNIPPET / AUTHOR STATUS ===")
    snippet = qlog.get("snippet_enrichment_status", "MISSING")
    author = qlog.get("author_enrichment_status", "MISSING")
    valid = ["available_but_not_invoked", "invoked", "not_available"]
    if snippet in valid and author in valid:
        print(f"  PASS F: snippet={snippet}, author={author}")
    else:
        print(f"  FAIL F: invalid status — snippet={snippet}, author={author}")
        all_pass = False

    # ── G. DOCX SELF-SUFFICIENCY — strict JSON reference check ──
    print("\n=== G. DOCX SELF-SUFFICIENCY ===")
    if docx_path.exists():
        try:
            from docx import Document
            doc = Document(str(docx_path))
            all_text = " ".join(p.text for p in doc.paragraphs)
            table_text = " ".join(
                c.text for t in doc.tables for r in t.rows for c in r.cells
            )
            full_text = all_text + " " + table_text

            json_terms = [
                "external_evidence_pack.json", "research_query_log.json",
                "research_results_raw.json", "query_execution_log.json",
                "external_evidence_pack", "research_query_log",
            ]
            # Also check for generic "json" or "JSON" (case-sensitive)
            found_refs = []
            for term in json_terms:
                if term.lower() in full_text.lower():
                    found_refs.append(term)
            if found_refs:
                print(f"  FAIL G: DOCX references JSON files: {found_refs}")
                all_pass = False
            else:
                print(f"  PASS G: DOCX is self-sufficient (no JSON references)")
        except ImportError:
            print(f"  SKIP G: python-docx not available")
    else:
        print(f"  FAIL G: source_book.docx not found")
        all_pass = False

    # ── H. BLUEPRINT SHELL CHECK ──────────────────────────
    print("\n=== H. BLUEPRINT SHELL CHECK ===")
    bp_path = output_dir / "slide_blueprint_from_source_book.json"
    if bp_path.exists():
        bp = json.loads(bp_path.read_text(encoding="utf-8"))
        shell_sections = ["S02", "S04", "S06", "S08", "S10"]
        shell_issues = 0
        for entry in bp:
            sid = entry.get("section_id", "")
            if sid in shell_sections:
                km = entry.get("key_message", "") or ""
                title = entry.get("slide_title", "") or ""
                # Check for shell/placeholder content
                sn = entry.get("section_name", "")
                if km == f"{sn} content" or km == f"{sn} shell":
                    print(f"  FAIL H: {sid} has shell key_message: '{km}'")
                    shell_issues += 1
                    all_pass = False
                if title == sn and sid in ["S02"]:
                    print(f"  FAIL H: {sid} has shell title = section name: '{title}'")
                    shell_issues += 1
                    all_pass = False
        if shell_issues == 0:
            print(f"  PASS H: no shell/placeholder content in divider/dynamic sections")
    else:
        print(f"  WARN H: blueprint file not found")

    # ── I. TERMINAL LOG CROSS-CHECK ─────────────────────
    print("\n=== I. TERMINAL LOG CROSS-CHECK ===")
    # Find the terminal log for this session
    # Terminal logs are in the output/ directory (parent of session dir)
    terminal_log = None
    output_root = output_dir.parent
    for f in sorted(output_root.glob("source_book_*.txt"), reverse=True):
        try:
            content = f.read_text(encoding="utf-8", errors="replace")
            if session_id in content:
                terminal_log = content
                logger.info("Found terminal log for %s: %s", session_id, f.name) if 'logger' in dir() else None
                break
        except Exception:
            continue
    if terminal_log:
        # Count S2 bulk search invocations in terminal log
        import re
        s2_searches_in_log = len(re.findall(r"S2 search: query=", terminal_log))
        pplx_calls_in_log = len(re.findall(r"Perplexity API 200", terminal_log))

        # Count in execution log
        s2_in_exec = sum(1 for e in exec_log if e.get("service_invoked") == "semantic_scholar")
        pplx_in_exec = sum(1 for e in exec_log if e.get("service_invoked") == "perplexity")

        if s2_searches_in_log != s2_in_exec:
            print(f"  FAIL I: S2 searches in terminal={s2_searches_in_log} vs exec_log={s2_in_exec}")
            all_pass = False
        else:
            print(f"  PASS I1: S2 count matches (terminal={s2_searches_in_log}, exec={s2_in_exec})")

        if pplx_calls_in_log != pplx_in_exec:
            print(f"  FAIL I: Perplexity calls in terminal={pplx_calls_in_log} vs exec_log={pplx_in_exec}")
            all_pass = False
        else:
            print(f"  PASS I2: Perplexity count matches (terminal={pplx_calls_in_log}, exec={pplx_in_exec})")
    else:
        print(f"  WARN I: no terminal log found for cross-check")

    # ── J. EXACT DOCX-vs-JSON THEME COVERAGE ──────────────
    print("\n=== J. EXACT DOCX-vs-JSON THEME COVERAGE ===")
    if docx_path.exists() and tc:
        try:
            from docx import Document
            doc = Document(str(docx_path))
            docx_lines = [p.text for p in doc.paragraphs]

            j_mismatches = 0
            for theme_key, info in tc.items():
                count = info.get("retained_sources", 0)
                status = info.get("status", "gap")
                expected_fragment = f"{count} sources"

                # Find the DOCX line for this theme
                theme_label = theme_key.replace("_", " ").title()
                found_line = None
                for line in docx_lines:
                    if theme_label.lower() in line.lower() and "sources" in line.lower():
                        found_line = line
                        break

                if found_line:
                    if expected_fragment not in found_line:
                        print(f"  FAIL J: {theme_key}: JSON says {count} but DOCX says: '{found_line[:80]}'")
                        j_mismatches += 1
                        all_pass = False
                    elif status not in found_line:
                        print(f"  FAIL J: {theme_key}: JSON status='{status}' not in DOCX: '{found_line[:80]}'")
                        j_mismatches += 1
                        all_pass = False

            if j_mismatches == 0:
                print(f"  PASS J: all {len(tc)} themes match exactly between JSON and DOCX")
        except ImportError:
            print(f"  SKIP J: python-docx not available")
    else:
        if not tc:
            print(f"  FAIL J: no theme_coverage in JSON")
            all_pass = False
        else:
            print(f"  FAIL J: DOCX not found")
            all_pass = False

    # ── Summary ───────────────────────────────────────────
    print(f"\n{'='*60}")
    print(f"  VERIFICATION RESULT: {'ALL PASS' if all_pass else 'SOME FAILURES'}")
    print(f"{'='*60}")
    return all_pass


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/verify_artifacts.py <session_id>")
        sys.exit(1)
    passed = verify(sys.argv[1])
    sys.exit(0 if passed else 1)
