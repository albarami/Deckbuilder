"""Artifact gates — context-specific usage rules for claims in the pipeline.

Every claim must pass the appropriate gate before appearing in an artifact.
The gates form a hierarchy from most permissive (internal_gap_appendix)
to most restrictive (client_proposal / proof_point).

Architecture invariant:
    absence of Engine 2 verification defaults to internal_unverified = blocked.
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING, Literal

from pydantic import Field, model_validator

from src.models.claim_provenance import ClaimProvenance, ClaimRegistry
from src.models.common import DeckForgeBaseModel

if TYPE_CHECKING:
    from src.models.conformance import ConformanceReport
    from src.models.source_book import SourceBookReview


# ── Context-specific gates ───────────────────────────────────────


def can_use_as_proof_point(claim: ClaimProvenance) -> bool:
    """Strictest gate: proof columns, capability evidence, slide proof_points.

    Only verified + permissioned claims pass.
    """
    if claim.verification_status == "forbidden":
        return False

    if claim.claim_kind == "rfp_fact":
        return claim.verification_status == "verified_from_rfp"

    if claim.claim_kind == "external_methodology":
        return (
            claim.verification_status == "externally_verified"
            and claim.relevance_class in ("direct_topic", "adjacent_domain")
            and claim.evidence_role == "methodology_support"
        )

    if claim.claim_kind == "internal_company_claim":
        if claim.verification_status != "internal_verified":
            return False
        if (
            claim.requires_client_naming_permission
            and claim.client_naming_permission is not True
        ):
            return False
        if (
            claim.requires_partner_naming_permission
            and claim.partner_naming_permission is not True
        ):
            return False
        if claim.scope_summary_allowed_for_proposal is False:
            return False
        return True

    # proposal_option, generated_inference: never proof points
    return False


def can_use_in_source_book_analysis(
    claim: ClaimProvenance,
    section_type: Literal[
        "client_facing_body",
        "internal_bid_notes",
        "internal_gap_appendix",
        "evidence_gap_register",
    ],
) -> bool:
    """Source book body: analysis, problem framing, methodology design.

    Internal sections allow everything. Client-facing body is selective:
    - rfp_facts: always
    - external_methodology: verified or partially_verified (with restrictions)
    - internal_company_claim: only if internal_verified
    - generated_inference: only if labelled + context-approved
    - proposal_option: never in client-facing body
    """
    if claim.verification_status == "forbidden":
        return False

    if section_type in ("internal_gap_appendix", "evidence_gap_register"):
        return True

    if section_type == "internal_bid_notes":
        return True  # everything except forbidden

    # client_facing_body rules:
    if claim.claim_kind == "rfp_fact":
        return True

    if claim.claim_kind == "external_methodology":
        if claim.verification_status == "externally_verified":
            return claim.relevance_class in (
                "direct_topic",
                "adjacent_domain",
                "analogical",
            )
        if claim.verification_status == "partially_verified":
            return (
                claim.relevance_class != "analogical"
                and claim.evidence_role
                in ("methodology_support", "risk_or_assumption_support")
            )
        return False

    if claim.claim_kind == "internal_company_claim":
        return claim.verification_status == "internal_verified"

    if claim.claim_kind == "generated_inference":
        return (
            claim.verification_status == "generated_inference"
            and claim.inference_label_present is True
            and "source_book_analysis" in claim.inference_allowed_context
        )

    if claim.claim_kind == "proposal_option":
        return False  # options go to internal bid strategy only

    return False


def can_use_in_slide_blueprint(claim: ClaimProvenance) -> bool:
    """Slide blueprints: more restrictive than source book.

    No unverified internal claims. No proposal options.
    Generated inferences blocked from slides entirely.
    """
    if claim.verification_status == "forbidden":
        return False
    if claim.claim_kind in (
        "rfp_fact",
        "external_methodology",
        "internal_company_claim",
    ):
        return can_use_as_proof_point(claim)
    return False  # generated_inference and proposal_option blocked


def can_use_in_speaker_notes(claim: ClaimProvenance) -> bool:
    """Speaker notes: allows labelled generated inferences."""
    return (
        claim.claim_kind == "generated_inference"
        and claim.inference_label_present is True
        and "speaker_notes" in claim.inference_allowed_context
    )


def can_use_in_client_proposal(claim: ClaimProvenance) -> bool:
    """Final proposal: strictest external-facing gate."""
    return can_use_as_proof_point(claim)


def can_use_in_internal_gap_appendix(claim: ClaimProvenance) -> bool:
    """Internal appendix: everything visible for review."""
    return True


# ── Usage normalization ──────────────────────────────────────────


def normalize_usage_allowed(claim: ClaimProvenance) -> list[str]:
    """Compute actual allowed usage from verification status.

    usage_allowed is subordinate to verification_status — this function
    overrides manually-set usage when the claim is unverified or forbidden.
    Does NOT touch requested_external_contexts.
    """
    if claim.verification_status in ("internal_unverified", "forbidden"):
        return ["internal_gap_appendix"]
    if can_use_as_proof_point(claim):
        return claim.usage_allowed
    return ["internal_gap_appendix"]


# ── Artifact Section ─────────────────────────────────────────────


class ArtifactSection(DeckForgeBaseModel):
    """A section of a generated artifact, tagged by type for scanning."""

    section_path: str
    section_type: Literal[
        "client_facing_body",
        "proof_column",
        "slide_body",
        "slide_proof_points",
        "speaker_notes",
        "internal_gap_appendix",
        "internal_bid_notes",
        "evidence_ledger",
        "drafting_notes",
    ]
    text: str


# ── Forbidden Leakage Scanner ────────────────────────────────────


FORBIDDEN_ID_PATTERNS = [
    r"\bPRJ-\d+\b",
    r"\bCLI-\d+\b",
    r"\bCLM-\d+\b",
    r"INTERNAL_PROOF_PLACEHOLDER",
    r"Engine\s*2\b",           # Generic Engine 2 references (covers "ENGINE 2 REQUIRED",
                               # "Engine 2 must retrieve", "Engine 2 action", "من Engine 2")
    r"إثبات داخلي مطلوب",
    r"قيد الاستكمال من سجلات",
]

FORBIDDEN_SEMANTIC_PHRASES = [
    "خبرة موثقة في العمل مع سدايا",
    "تعاون موثق مع اليونسكو",
    "مشروع سابق مع سدايا واليونسكو",
    "documented experience with SDAIA",
    "documented collaboration with UNESCO",
]

INTERNAL_ONLY_SECTION_TYPES = {
    "internal_gap_appendix",
    "drafting_notes",
    "evidence_ledger",
    "internal_bid_notes",
}


class ForbiddenLeakageViolation(DeckForgeBaseModel):
    """A single violation found by the forbidden scanner."""

    pattern: str
    matched_text: str
    location: str
    section_type: str
    severity: Literal["critical", "major"] = "critical"
    linked_claim_id: str | None = None


def scan_for_forbidden_leakage(
    section: ArtifactSection,
    registry: ClaimRegistry | None = None,
) -> list[ForbiddenLeakageViolation]:
    """Scan a section for forbidden internal-claim leakage.

    Returns empty list if section is an internal-only type.
    """
    if section.section_type in INTERNAL_ONLY_SECTION_TYPES:
        return []

    violations: list[ForbiddenLeakageViolation] = []

    for pattern in FORBIDDEN_ID_PATTERNS:
        for match in re.finditer(pattern, section.text, re.IGNORECASE):
            violations.append(ForbiddenLeakageViolation(
                pattern=pattern,
                matched_text=match.group(),
                location=section.section_path,
                section_type=section.section_type,
            ))

    for phrase in FORBIDDEN_SEMANTIC_PHRASES:
        if phrase in section.text:
            # Registry-aware: check if linked claim is actually verified
            if registry:
                linked = registry.resolve_proof_point(phrase)
                if linked and can_use_as_proof_point(linked):
                    continue  # verified + permissioned: allow
            violations.append(ForbiddenLeakageViolation(
                pattern=f"semantic:{phrase}",
                matched_text=phrase,
                location=section.section_path,
                section_type=section.section_type,
            ))

    return violations


# ── Source Book → ArtifactSection adapter ────────────────────────


def render_source_book_sections(source_book: object) -> list[ArtifactSection]:
    """Render a SourceBook into typed ArtifactSection records for scanning.

    Sections are tagged so the forbidden-leakage scanner can apply the
    correct policy: client_facing_body / proof_column / slide_body /
    slide_proof_points get scanned; evidence_ledger and the internal
    appendices are exempt because PRJ/CLI/CLM identifiers are expected
    there.

    Source: design doc Section 3 (forbidden scanner section-type table).
    """
    sections: list[ArtifactSection] = []

    rfp_int = getattr(source_book, "rfp_interpretation", None)
    if rfp_int is not None:
        for field in (
            "objective_and_scope",
            "constraints_and_compliance",
            "unstated_evaluator_priorities",
            "probable_scoring_logic",
        ):
            text = getattr(rfp_int, field, "") or ""
            if text:
                sections.append(ArtifactSection(
                    section_path=f"rfp_interpretation/{field}",
                    section_type="client_facing_body",
                    text=text,
                ))
        for i, item in enumerate(getattr(rfp_int, "key_compliance_requirements", []) or []):
            sections.append(ArtifactSection(
                section_path=f"rfp_interpretation/key_compliance[{i}]",
                section_type="client_facing_body",
                text=str(item),
            ))

    framing = getattr(source_book, "client_problem_framing", None)
    if framing is not None:
        for field in (
            "current_state_challenge",
            "why_it_matters_now",
        ):
            text = getattr(framing, field, "") or ""
            if text:
                sections.append(ArtifactSection(
                    section_path=f"client_problem_framing/{field}",
                    section_type="client_facing_body",
                    text=text,
                ))

    why_sg = getattr(source_book, "why_strategic_gears", None)
    if why_sg is not None:
        for i, cm in enumerate(getattr(why_sg, "capability_mapping", []) or []):
            text_bits = []
            for field in ("capability", "approach", "evidence_summary"):
                val = getattr(cm, field, "") or ""
                if val:
                    text_bits.append(str(val))
            if text_bits:
                sections.append(ArtifactSection(
                    section_path=f"why_strategic_gears/capability_mapping[{i}]",
                    section_type="proof_column",
                    text=" ".join(text_bits),
                ))
        for i, cert in enumerate(getattr(why_sg, "certifications_and_compliance", []) or []):
            sections.append(ArtifactSection(
                section_path=f"why_strategic_gears/certifications[{i}]",
                section_type="proof_column",
                text=str(cert),
            ))
        for i, pe in enumerate(getattr(why_sg, "project_experience", []) or []):
            text_bits = []
            for field in ("project_name", "client", "scope_summary", "outcome_summary"):
                val = getattr(pe, field, "") or ""
                if val:
                    text_bits.append(str(val))
            if text_bits:
                sections.append(ArtifactSection(
                    section_path=f"why_strategic_gears/project_experience[{i}]",
                    section_type="proof_column",
                    text=" ".join(text_bits),
                ))

    proposed = getattr(source_book, "proposed_solution", None)
    if proposed is not None:
        for field in (
            "methodology_overview",
            "governance_framework",
            "timeline_logic",
            "value_case_and_differentiation",
        ):
            text = getattr(proposed, field, "") or ""
            if text:
                sections.append(ArtifactSection(
                    section_path=f"proposed_solution/{field}",
                    section_type="client_facing_body",
                    text=text,
                ))

    for i, blueprint in enumerate(getattr(source_book, "slide_blueprints", []) or []):
        body_bits = []
        for field in ("title", "key_message", "purpose", "visual_guidance"):
            val = getattr(blueprint, field, "") or ""
            if val:
                body_bits.append(str(val))
        for bullet in getattr(blueprint, "bullet_logic", []) or []:
            body_bits.append(str(bullet))
        if body_bits:
            sections.append(ArtifactSection(
                section_path=f"slide_blueprints[{i}]/body",
                section_type="slide_body",
                text=" ".join(body_bits),
            ))
        proof_bits = []
        for proof in getattr(blueprint, "proof_points", []) or []:
            proof_bits.append(str(proof))
        for must in getattr(blueprint, "must_have_evidence", []) or []:
            proof_bits.append(str(must))
        if proof_bits:
            sections.append(ArtifactSection(
                section_path=f"slide_blueprints[{i}]/proof_points",
                section_type="slide_proof_points",
                text=" ".join(proof_bits),
            ))

    # Evidence ledger is internal-only — we still emit a section for it so
    # downstream callers (e.g. final_artifact_gate) see the surface, but
    # the scanner skips sections of type evidence_ledger.
    ledger = getattr(source_book, "evidence_ledger", None)
    if ledger is not None:
        for i, entry in enumerate(getattr(ledger, "entries", []) or []):
            text_bits = []
            for field in ("claim_id", "claim_text", "source_reference"):
                val = getattr(entry, field, "") or ""
                if val:
                    text_bits.append(str(val))
            if text_bits:
                sections.append(ArtifactSection(
                    section_path=f"evidence_ledger/entries[{i}]",
                    section_type="evidence_ledger",
                    text=" ".join(text_bits),
                ))

    return sections


# ── Slide proof-point gating (Slice 2.3) ─────────────────────────


class ProofPointGatingViolation(DeckForgeBaseModel):
    """One slide proof_point that failed registry resolution or proof gating."""

    slide_number: int
    proof_point: str
    reason: Literal["unresolved_in_registry", "not_a_proof_point"]
    claim_id: str = ""
    verification_status: str = ""


def gate_slide_proof_points(
    blueprints: list[object],
    registry: ClaimRegistry,
) -> tuple[list[object], list[ProofPointGatingViolation]]:
    """Resolve every proof_point + must_have_evidence against the
    ClaimRegistry and drop anything that doesn't pass
    ``can_use_as_proof_point``.

    Returns ``(gated_blueprints, violations)``. Each violation records the
    slide_number, the offending proof string, the reason, and (when the
    claim was found) the claim_id + verification_status.

    PRJ-/CLI-/CLM- identifiers that don't resolve to a registered claim
    fail with ``unresolved_in_registry``. Resolved claims that are
    unverified (or external methodology with the wrong evidence_role,
    proposal options, etc.) fail with ``not_a_proof_point``.
    """
    gated: list[object] = []
    violations: list[ProofPointGatingViolation] = []

    for bp in blueprints:
        slide_number = int(getattr(bp, "slide_number", 0) or 0)

        def _gate_one(candidate: str) -> bool:
            claim = registry.resolve_proof_point(candidate)
            if claim is None:
                violations.append(ProofPointGatingViolation(
                    slide_number=slide_number,
                    proof_point=candidate,
                    reason="unresolved_in_registry",
                ))
                return False
            if not can_use_as_proof_point(claim):
                violations.append(ProofPointGatingViolation(
                    slide_number=slide_number,
                    proof_point=candidate,
                    reason="not_a_proof_point",
                    claim_id=claim.claim_id,
                    verification_status=claim.verification_status,
                ))
                return False
            return True

        kept_proofs = [p for p in (getattr(bp, "proof_points", []) or []) if _gate_one(p)]
        kept_must = [m for m in (getattr(bp, "must_have_evidence", []) or []) if _gate_one(m)]

        # Use model_copy so the entry's own validators (e.g. auto-restore
        # proof_points from must_have_evidence) re-run against the gated lists.
        gated.append(bp.model_copy(update={
            "proof_points": kept_proofs,
            "must_have_evidence": kept_must,
        }))

    return gated, violations


# ── Evidence Coverage ────────────────────────────────────────────


class EvidenceCoverageRequirement(DeckForgeBaseModel):
    """Coverage requirement for one evidence topic."""

    topic: str
    minimum_direct_sources: int
    found_direct: int = 0
    found_adjacent: int = 0
    found_analogical: int = 0
    status: Literal["met", "not_met"] = "not_met"
    missing_reason: str = ""

    @model_validator(mode="after")
    def compute_status(self) -> "EvidenceCoverageRequirement":
        object.__setattr__(
            self, "status",
            "met" if self.found_direct >= self.minimum_direct_sources else "not_met",
        )
        return self


class EvidenceCoverageReport(DeckForgeBaseModel):
    """Aggregated evidence coverage across all required topics."""

    requirements: list[EvidenceCoverageRequirement] = Field(default_factory=list)
    status: Literal["pass", "fail"] = "fail"

    @model_validator(mode="after")
    def compute_status(self) -> "EvidenceCoverageReport":
        all_met = all(r.status == "met" for r in self.requirements)
        object.__setattr__(self, "status", "pass" if all_met else "fail")
        return self


# ── Coverage builder (Slice 3.2) ────────────────────────────────


class CoverageTopic(DeckForgeBaseModel):
    """Per-topic spec used to drive build_evidence_coverage_report.

    keywords are case-insensitive substrings searched against each
    external_methodology claim's text. minimum_direct_sources is the
    bar a topic must hit on direct_topic claims alone — adjacent and
    analogical sources do not satisfy it (Slice 3 acceptance #3).
    """

    name: str
    keywords: list[str]
    minimum_direct_sources: int = 1


def build_evidence_coverage_report(
    registry: "ClaimRegistry",
    topics: list[CoverageTopic],
) -> EvidenceCoverageReport:
    """Count external_methodology claims per topic and bucket by
    relevance_class. Only direct_topic counts toward
    minimum_direct_sources."""
    requirements: list[EvidenceCoverageRequirement] = []
    methodology_claims = list(registry.external_methodology)

    for topic in topics:
        direct = adjacent = analogical = 0
        for claim in methodology_claims:
            text = (claim.text or "").lower()
            if not any(kw.lower() in text for kw in topic.keywords if kw):
                continue
            if claim.relevance_class == "direct_topic":
                direct += 1
            elif claim.relevance_class == "adjacent_domain":
                adjacent += 1
            elif claim.relevance_class == "analogical":
                analogical += 1
        requirements.append(EvidenceCoverageRequirement(
            topic=topic.name,
            minimum_direct_sources=topic.minimum_direct_sources,
            found_direct=direct,
            found_adjacent=adjacent,
            found_analogical=analogical,
        ))

    return EvidenceCoverageReport(requirements=requirements)


# ── Gate Failure ─────────────────────────────────────────────────


class GateFailure(DeckForgeBaseModel):
    """A single failure from the final gate."""

    code: str
    severity: Literal["critical", "major", "minor"]
    message: str
    affected_artifact: str | None = None
    claim_id: str | None = None


class ArtifactGateDecision(DeckForgeBaseModel):
    """Decision from the final artifact gate."""

    decision: Literal["approve", "reject"]
    proposal_ready: bool
    deck_generation_allowed: bool
    artifact_label: str
    failures: list[GateFailure] = Field(default_factory=list)


# ── Final Artifact Gate ──────────────────────────────────────────


def scan_docx_for_forbidden_leakage(docx_path: str) -> list[dict]:
    """Scan an exported DOCX file for forbidden internal-proof patterns.

    Checks both FORBIDDEN_ID_PATTERNS (regex) and FORBIDDEN_SEMANTIC_PHRASES
    (substring). Returns a list of finding dicts, each with pattern,
    matched_text, and source fields.

    Raises on file-not-found or parse errors — caller must handle.
    """
    from docx import Document as DocxDoc

    doc = DocxDoc(docx_path)
    all_text: list[str] = []
    for p in doc.paragraphs:
        all_text.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_text.append(p.text)
    docx_text = "\n".join(all_text)

    findings: list[dict] = []

    # ID patterns (regex)
    for pattern_str in FORBIDDEN_ID_PATTERNS:
        for m in re.finditer(pattern_str, docx_text, re.IGNORECASE):
            findings.append({
                "pattern": pattern_str,
                "matched_text": m.group(),
                "source": "post_export_docx_scan",
            })

    # Semantic phrases (substring)
    for phrase in FORBIDDEN_SEMANTIC_PHRASES:
        if phrase in docx_text:
            findings.append({
                "pattern": f"semantic:{phrase}",
                "matched_text": phrase,
                "source": "post_export_docx_scan",
            })

    return findings


def final_artifact_gate(
    conformance_report: "ConformanceReport",
    reviewer_score: "SourceBookReview",
    evidence_coverage: EvidenceCoverageReport,
    forbidden_scan: list[ForbiddenLeakageViolation],
    claim_registry: ClaimRegistry,
    rendered_sections: list[ArtifactSection],
) -> ArtifactGateDecision:
    """Final pipeline checkpoint. ALL gates must pass for proposal_ready.

    Checks: conformance, forbidden claims, reviewer threshold,
    evidence coverage, forbidden leakage, rendered artifact scanning.
    """
    failures: list[GateFailure] = []

    if conformance_report.conformance_status != "pass":
        failures.append(GateFailure(
            code="CONFORMANCE_FAIL",
            severity="critical",
            message=f"conformance_status={conformance_report.conformance_status}",
        ))

    forbidden_count = getattr(conformance_report, "conformance_forbidden_claims", 0)
    if forbidden_count and forbidden_count > 0:
        failures.append(GateFailure(
            code="FORBIDDEN_CLAIMS",
            severity="critical",
            message=f"forbidden_claims={forbidden_count}",
        ))

    if not reviewer_score.pass_threshold_met:
        failures.append(GateFailure(
            code="REVIEWER_THRESHOLD",
            severity="major",
            message=f"reviewer_threshold_met=False (score={reviewer_score.overall_score})",
        ))

    if evidence_coverage.status != "pass":
        failures.append(GateFailure(
            code="EVIDENCE_COVERAGE",
            severity="critical",
            message=f"evidence_coverage={evidence_coverage.status}",
        ))

    if len(forbidden_scan) > 0:
        failures.append(GateFailure(
            code="FORBIDDEN_LEAKAGE",
            severity="critical",
            message=f"forbidden_leakage={len(forbidden_scan)} violations",
        ))

    # Scan rendered artifacts for forbidden leakage
    for section in rendered_sections:
        section_violations = scan_for_forbidden_leakage(section, claim_registry)
        for v in section_violations:
            failures.append(GateFailure(
                code="RENDERED_LEAKAGE",
                severity="critical",
                message=f"{v.matched_text} in {v.location}",
                claim_id=v.linked_claim_id,
            ))

    if any(f.severity in ("critical", "major") for f in failures):
        return ArtifactGateDecision(
            decision="reject",
            proposal_ready=False,
            deck_generation_allowed=False,
            artifact_label="DRAFT — NOT PROPOSAL READY",
            failures=failures,
        )

    return ArtifactGateDecision(
        decision="approve",
        proposal_ready=True,
        deck_generation_allowed=True,
        artifact_label="PROPOSAL READY",
        failures=failures,
    )
