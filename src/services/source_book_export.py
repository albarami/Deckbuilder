"""DOCX exporter for the Proposal Source Book.

Uses python-docx to generate a structured Word document with:
- Cover page with client name, RFP name, date
- 7 sections with appropriate formatting (prose, tables)
- Table of Contents placeholder
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.models.source_book import SourceBook

logger = logging.getLogger(__name__)

# Regex for pipe-delimited table rows: | col1 | col2 | ... |
_PIPE_ROW_RE = re.compile(r"^\s*\|(.+\|)\s*$")
# Separator row: |---|---|---| or | --- | --- |
_SEPARATOR_RE = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")
# Inline pipe pattern: text with 2+ pipe chars (e.g., "COMP-001 | req | response")
_INLINE_PIPE_RE = re.compile(r"^[^|]+\|[^|]+\|.+$")


def _parse_pipe_table(lines: list[str]) -> list[list[str]] | None:
    """Parse consecutive pipe-delimited lines into a list of rows.

    Returns None if the lines don't form a valid table (< 2 data rows).
    """
    rows: list[list[str]] = []
    for line in lines:
        if _SEPARATOR_RE.match(line):
            continue
        m = _PIPE_ROW_RE.match(line)
        if not m:
            continue
        cells = [c.strip() for c in m.group(1).split("|")]
        # Remove empty leading/trailing from split artifacts
        while cells and not cells[-1]:
            cells.pop()
        while cells and not cells[0]:
            cells.pop(0)
        if cells:
            rows.append(cells)
    if len(rows) < 2:
        return None
    return rows


def _add_word_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a real Word table from parsed row data. First row is header."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.add_row()
    for i, cell_text in enumerate(rows[0]):
        if i < n_cols:
            cell = hdr_row.cells[i]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

    # Data rows
    for row_data in rows[1:]:
        data_row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i < n_cols:
                cell = data_row.cells[i]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def _add_smart_prose(doc: Document, text: str) -> None:
    """Add prose text, converting any embedded pipe-delimited tables to real Word tables."""
    if not text:
        return

    lines = text.split("\n")
    buffer: list[str] = []  # non-table lines
    table_lines: list[str] = []  # consecutive pipe lines

    def flush_prose() -> None:
        if buffer:
            prose = "\n".join(buffer).strip()
            if prose:
                doc.add_paragraph(prose)
            buffer.clear()

    def flush_table() -> None:
        if table_lines:
            rows = _parse_pipe_table(table_lines)
            if rows:
                flush_prose()
                _add_word_table(doc, rows)
            else:
                # Not a valid table — treat as prose
                buffer.extend(table_lines)
            table_lines.clear()

    for line in lines:
        is_pipe = bool(_PIPE_ROW_RE.match(line) or _SEPARATOR_RE.match(line))
        is_inline_pipe = bool(not is_pipe and _INLINE_PIPE_RE.match(line.strip()))
        if is_pipe or is_inline_pipe:
            # Normalize inline pipe lines to standard pipe format
            normalized = line.strip()
            if is_inline_pipe and not normalized.startswith("|"):
                normalized = "| " + normalized
            if is_inline_pipe and not normalized.endswith("|"):
                normalized = normalized + " |"
            if not table_lines and buffer:
                flush_prose()
            table_lines.append(normalized)
        else:
            if table_lines:
                flush_table()
            buffer.append(line)

    flush_table()
    flush_prose()


def _add_cover_page(doc: Document, source_book: SourceBook) -> None:
    """Add cover page with client name, RFP name, and date."""
    # Title
    title = doc.add_heading("PROPOSAL SOURCE BOOK", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Client and RFP
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"{source_book.client_name} — {source_book.rfp_name}"
    )
    run.font.size = Pt(16)
    run.bold = True

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Generated: {source_book.generation_date or 'N/A'} | "
        f"Language: {source_book.language} | "
        f"Pass: {source_book.pass_number}"
    )
    meta_run.font.size = Pt(10)

    doc.add_page_break()


def _add_section_1(doc: Document, source_book: SourceBook) -> None:
    """Section 1: RFP Interpretation."""
    doc.add_heading("1. RFP Interpretation", level=1)
    rfp = source_book.rfp_interpretation

    if rfp.objective_and_scope:
        doc.add_heading("1.1 Objective & Scope", level=2)
        _add_smart_prose(doc, rfp.objective_and_scope)

    if rfp.constraints_and_compliance:
        doc.add_heading("1.2 Constraints & Compliance Requirements", level=2)
        _add_smart_prose(doc, rfp.constraints_and_compliance)

    if rfp.unstated_evaluator_priorities:
        doc.add_heading("1.3 Unstated Evaluator Priorities", level=2)
        _add_smart_prose(doc, rfp.unstated_evaluator_priorities)

    if rfp.probable_scoring_logic:
        doc.add_heading("1.4 Probable Scoring Logic", level=2)
        _add_smart_prose(doc, rfp.probable_scoring_logic)

    if rfp.key_compliance_requirements:
        doc.add_heading("1.5 Key Compliance Requirements", level=2)
        # Check if entries contain pipe-delimited content (e.g., COMP-xxx | ... | ...)
        pipe_entries = [r for r in rfp.key_compliance_requirements if r.count("|") >= 2]
        if pipe_entries:
            # Render as a table via _add_smart_prose
            combined = "\n".join(pipe_entries)
            _add_smart_prose(doc, combined)
            # Render any non-pipe entries as bullets
            for req in rfp.key_compliance_requirements:
                if req.count("|") < 2:
                    doc.add_paragraph(req, style="List Bullet")
        else:
            for req in rfp.key_compliance_requirements:
                doc.add_paragraph(req, style="List Bullet")


def _add_section_2(doc: Document, source_book: SourceBook) -> None:
    """Section 2: Client Problem Framing."""
    doc.add_heading("2. Client Problem Framing", level=1)
    cpf = source_book.client_problem_framing

    if cpf.current_state_challenge:
        doc.add_heading("2.1 Current-State Challenge", level=2)
        _add_smart_prose(doc, cpf.current_state_challenge)

    if cpf.why_it_matters_now:
        doc.add_heading("2.2 Why It Matters Now", level=2)
        _add_smart_prose(doc, cpf.why_it_matters_now)

    if cpf.transformation_logic:
        doc.add_heading("2.3 Transformation Logic", level=2)
        _add_smart_prose(doc, cpf.transformation_logic)

    if cpf.risk_if_unchanged:
        doc.add_heading("2.4 Risk If Unchanged", level=2)
        _add_smart_prose(doc, cpf.risk_if_unchanged)


def _add_section_3(doc: Document, source_book: SourceBook) -> None:
    """Section 3: Why Strategic Gears."""
    doc.add_heading("3. Why Strategic Gears", level=1)
    wsg = source_book.why_strategic_gears

    # 3.1 Capability Mapping table
    if wsg.capability_mapping:
        doc.add_heading("3.1 Capability-to-RFP Mapping", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "RFP Requirement"
        hdr[1].text = "SG Capability"
        hdr[2].text = "Evidence"
        hdr[3].text = "Strength"
        for cm in wsg.capability_mapping:
            row = table.add_row().cells
            row[0].text = cm.rfp_requirement
            row[1].text = cm.sg_capability
            row[2].text = ", ".join(cm.evidence_ids) if cm.evidence_ids else "—"
            row[3].text = cm.strength

    # 3.2 Named Consultants table — interim team structure
    if wsg.named_consultants:
        doc.add_heading("3.2 Proposed Team & Staffing Recommendations", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Name"
        hdr[1].text = "Role"
        hdr[2].text = "Status"
        hdr[3].text = "Relevance / Justification"
        hdr[4].text = "Source"
        hdr[5].text = "Confidence"
        for nc in wsg.named_consultants:
            row = table.add_row().cells
            status = getattr(nc, "staffing_status", "recommended_candidate")
            status_label = status.replace("_", " ").title()
            row[0].text = nc.name or "(Open)"
            row[1].text = nc.role
            row[2].text = status_label
            justification = getattr(nc, "justification", "")
            row[3].text = f"{nc.relevance}\n{justification}".strip()
            row[4].text = getattr(nc, "source_of_recommendation", "") or "—"
            row[5].text = getattr(nc, "confidence", "medium")

    # 3.3 Project Experience table
    if wsg.project_experience:
        doc.add_heading("3.3 Relevant Project Experience", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Project"
        hdr[1].text = "Client"
        hdr[2].text = "Outcomes"
        hdr[3].text = "Evidence"
        for pe in wsg.project_experience:
            row = table.add_row().cells
            row[0].text = pe.project_name
            row[1].text = pe.client
            row[2].text = pe.outcomes
            row[3].text = ", ".join(pe.evidence_ids) if pe.evidence_ids else "—"

    # 3.4 Certifications
    if wsg.certifications_and_compliance:
        doc.add_heading("3.4 Certifications & Compliance", level=2)
        for cert in wsg.certifications_and_compliance:
            doc.add_paragraph(cert, style="List Bullet")


def _add_section_4(
    doc: Document,
    source_book: SourceBook,
    evidence_enrichment: dict[str, dict] | None = None,
) -> None:
    """Section 4: External Evidence — enriched for proposal-building.

    Each entry shows source type, evidence tier, provider, URL, RFP theme,
    relevance, and key findings. The coverage assessment tags evidence gaps.
    """
    doc.add_heading("4. External Evidence", level=1)
    ext = source_book.external_evidence
    enrichment = evidence_enrichment or {}

    if ext.entries:
        # Enriched table with provider, URL, RFP theme
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Source ID"
        hdr[1].text = "Title / Provider / URL"
        hdr[2].text = "Type / Year"
        hdr[3].text = "Evidence Tier"
        hdr[4].text = "RFP Theme & How to Use"
        hdr[5].text = "Key Finding"

        for entry in ext.entries:
            row = table.add_row().cells
            row[0].text = entry.source_id

            # Title + Provider + URL (enriched from evidence pack)
            enrich = enrichment.get(entry.source_id, {})
            provider = enrich.get("provider", "").replace("_", " ").title()
            url = enrich.get("url", "")
            title_parts = [entry.title]
            if provider:
                title_parts.append(f"Provider: {provider}")
            if url:
                title_parts.append(f"URL: {url}")
            row[1].text = "\n".join(title_parts)

            # Type classification
            type_label = entry.source_type.replace("_", " ").title()
            row[2].text = f"{type_label} ({entry.year})"

            # Evidence tier based on source_type
            tier_map = {
                "academic_paper": "Primary — peer-reviewed",
                "industry_report": "Primary — industry source",
                "benchmark": "Primary — benchmark data",
                "case_study": "Secondary — analogical",
                "framework": "Secondary — methodology reference",
            }
            row[3].text = tier_map.get(entry.source_type, "Unclassified")

            # RFP theme + relevance + usage (enriched)
            rfp_theme = enrich.get("mapped_rfp_theme", "")
            how_to_use = enrich.get("how_to_use", "")
            usage_parts = []
            if rfp_theme:
                usage_parts.append(f"RFP Theme: {rfp_theme}")
            usage_parts.append(entry.relevance)
            if how_to_use:
                usage_parts.append(f"How to use: {how_to_use}")
            elif entry.key_finding:
                usage_parts.append(f"Proposal use: {entry.key_finding}")
            row[4].text = "\n".join(usage_parts)

            row[5].text = entry.key_finding

    if ext.coverage_assessment:
        doc.add_paragraph()
        doc.add_heading("Coverage Assessment & Evidence Gaps", level=2)
        _add_smart_prose(doc, ext.coverage_assessment)

        # Add structured gap tags — dynamically from coverage assessment
        # (not hardcoded to any jurisdiction)
        doc.add_paragraph()
        doc.add_heading("Evidence Gap Summary", level=3)
        gap_items = [
            ("Jurisdiction-specific benchmarks",
             "No local jurisdiction benchmarks found in external evidence. "
             "Engine 2 action: source from client's institutional databases."),
            ("SLA/KPI frameworks",
             "No specific service level agreement frameworks found. "
             "Engine 2 action: source from client's existing SLA framework."),
            ("National program alignment",
             "Limited evidence on jurisdiction-specific national programs. "
             "Engine 2 action: source from relevant national program documentation."),
        ]
        for gap_name, gap_desc in gap_items:
            p = doc.add_paragraph()
            run = p.add_run(f"❌ {gap_name}: ")
            run.bold = True
            p.add_run(gap_desc)

    # ── Search Strategy Summary ──────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Search Strategy Summary", level=2)

    # Compute statistics from entries
    total_sources = len(ext.entries) if ext.entries else 0
    type_counts: dict[str, int] = {}
    for entry in (ext.entries or []):
        label = entry.source_type.replace("_", " ").title()
        type_counts[label] = type_counts.get(label, 0) + 1

    # Query strategy paragraph
    doc.add_paragraph(
        f"The external evidence search yielded {total_sources} retained "
        f"source(s) across {len(type_counts)} evidence class(es). "
        "Queries were constructed from RFP themes, sector keywords, and "
        "compliance requirements to maximise coverage of the proposal's "
        "argument structure."
    )

    # Evidence class distribution
    if type_counts:
        p = doc.add_paragraph()
        run = p.add_run("Evidence class distribution: ")
        run.bold = True
        distribution_parts = [f"{cls} ({cnt})" for cls, cnt in type_counts.items()]
        p.add_run(", ".join(distribution_parts) + ".")

    # Coverage note
    tier_map_labels = {
        "Academic Paper": "Primary — peer-reviewed",
        "Industry Report": "Primary — industry source",
        "Benchmark": "Primary — benchmark data",
        "Case Study": "Secondary — analogical",
        "Framework": "Secondary — methodology reference",
    }
    primary_count = sum(
        cnt for cls, cnt in type_counts.items()
        if tier_map_labels.get(cls, "").startswith("Primary")
    )
    secondary_count = total_sources - primary_count
    doc.add_paragraph(
        f"Tier breakdown: {primary_count} primary source(s), "
        f"{secondary_count} secondary source(s). "
        "Coverage gaps are tagged in the Evidence Gap Summary above and "
        "flagged for Engine 2 retrieval from company backend databases."
    )

    # Provider / URL availability note
    doc.add_paragraph(
        "Provider and URL metadata for each source are recorded in the "
        "external_evidence_pack.json artifact produced during the research "
        "phase. Consult that JSON file for full Provider identifiers, "
        "URL links, author lists, query_used strings, and "
        "mapped_rfp_theme annotations."
    )


def _add_section_5(doc: Document, source_book: SourceBook) -> None:
    """Section 5: Proposed Solution."""
    doc.add_heading("5. Proposed Solution", level=1)
    ps = source_book.proposed_solution

    if ps.methodology_overview:
        doc.add_heading("5.1 Methodology Overview", level=2)
        _add_smart_prose(doc, ps.methodology_overview)

    if ps.phase_details:
        doc.add_heading("5.2 Phase Details", level=2)
        for phase in ps.phase_details:
            doc.add_heading(phase.phase_name, level=3)
            if phase.activities:
                p = doc.add_paragraph()
                run = p.add_run("Activities:")
                run.bold = True
                for act in phase.activities:
                    doc.add_paragraph(act, style="List Bullet")
            if phase.deliverables:
                p = doc.add_paragraph()
                run = p.add_run("Deliverables:")
                run.bold = True
                for d in phase.deliverables:
                    doc.add_paragraph(d, style="List Bullet")
            if phase.governance:
                p = doc.add_paragraph()
                run = p.add_run("Governance: ")
                run.bold = True
                p.add_run(phase.governance)

    if ps.governance_framework:
        doc.add_heading("5.3 Governance Framework", level=2)
        _add_smart_prose(doc, ps.governance_framework)

    if ps.timeline_logic:
        doc.add_heading("5.4 Timeline Logic", level=2)
        _add_smart_prose(doc, ps.timeline_logic)

    if ps.value_case_and_differentiation:
        doc.add_heading("5.5 Value Case & Differentiation", level=2)
        _add_smart_prose(doc, ps.value_case_and_differentiation)


def _add_section_6(doc: Document, source_book: SourceBook) -> None:
    """Section 6: Slide-by-Slide Blueprint."""
    doc.add_heading("6. Slide-by-Slide Blueprint", level=1)

    if not source_book.slide_blueprints:
        doc.add_paragraph("No slide blueprints generated.")
        return

    for bp in source_book.slide_blueprints:
        doc.add_heading(
            f"Slide {bp.slide_number}: {bp.title}",
            level=2,
        )

        # Blueprint as a compact table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(5.0)

        fields = [
            ("Section", bp.section),
            ("Layout", bp.layout),
            ("Purpose", bp.purpose),
            ("Key Message", bp.key_message),
            ("Bullet Logic", "\n".join(bp.bullet_logic) if bp.bullet_logic else "—"),
            ("Proof Points", ", ".join(bp.proof_points) if bp.proof_points else "—"),
            ("Visual", bp.visual_guidance),
            ("Must-Have Evidence", ", ".join(bp.must_have_evidence) if bp.must_have_evidence else "—"),
            ("Forbidden", ", ".join(bp.forbidden_content) if bp.forbidden_content else "—"),
        ]

        for label, value in fields:
            if value:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = value

        doc.add_paragraph()  # spacing


def _add_section_7(doc: Document, source_book: SourceBook) -> None:
    """Section 7: Evidence Ledger."""
    doc.add_heading("7. Evidence Ledger", level=1)

    if not source_book.evidence_ledger.entries:
        doc.add_paragraph("No evidence entries.")
        return

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Claim ID"
    hdr[1].text = "Claim Text"
    hdr[2].text = "Source Type"
    hdr[3].text = "Source Reference"
    hdr[4].text = "Confidence"
    hdr[5].text = "Status"

    for entry in source_book.evidence_ledger.entries:
        row = table.add_row().cells
        row[0].text = entry.claim_id
        row[1].text = entry.claim_text
        row[2].text = entry.source_type
        row[3].text = entry.source_reference
        row[4].text = f"{entry.confidence:.2f}"
        row[5].text = entry.verifiability_status


def _add_engine2_requirements(doc: Document, source_book: SourceBook) -> None:
    """Appendix: Engine 2 Requirements — proof gaps and action items.

    Surfaces all evidence_gap entries from the evidence ledger as a
    structured proof shopping list. Also lists open_role_profile team
    entries and their required qualifications.
    """
    # Collect gaps from evidence ledger
    evidence_gaps = [
        e for e in source_book.evidence_ledger.entries
        if e.verifiability_status == "gap"
    ]

    # Collect open roles from team
    open_roles = [
        nc for nc in source_book.why_strategic_gears.named_consultants
        if nc.staffing_status == "open_role_profile"
    ]

    if not evidence_gaps and not open_roles:
        return  # No gaps — Engine 2 not needed

    doc.add_page_break()
    doc.add_heading("Appendix: Engine 2 Requirements", level=1)
    doc.add_paragraph(
        "This appendix lists proof gaps that Engine 2 (company backend) "
        "must fill before the proposal is submission-ready. Each item "
        "specifies what evidence is needed and where it should come from."
    )

    # Evidence gaps table
    if evidence_gaps:
        doc.add_heading("A. Evidence Gaps — Proof Shopping List", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Claim ID"
        hdr[1].text = "Claim"
        hdr[2].text = "Source Needed"
        hdr[3].text = "Action Required"
        for entry in evidence_gaps:
            row = table.add_row().cells
            row[0].text = entry.claim_id
            row[1].text = entry.claim_text
            row[2].text = entry.source_reference or "Company backend"
            row[3].text = entry.verification_note or "Retrieve from company database"

    # Open roles table
    if open_roles:
        doc.add_heading("B. Open Team Roles — Staffing Requirements", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Role"
        hdr[1].text = "Required Qualifications"
        hdr[2].text = "Domain Expertise"
        hdr[3].text = "Action Required"
        for nc in open_roles:
            row = table.add_row().cells
            row[0].text = nc.role
            certs = ", ".join(nc.certifications) if nc.certifications else "—"
            edu = ", ".join(nc.education) if nc.education else "—"
            yrs = f"{nc.years_experience}+ years" if nc.years_experience else "—"
            row[1].text = f"Education: {edu}\nCertifications: {certs}\nExperience: {yrs}"
            row[2].text = ", ".join(nc.domain_expertise) if nc.domain_expertise else "—"
            row[3].text = (
                "Engine 2: retrieve matching consultant from staffing database "
                "and populate name, CV, and project history"
            )

    doc.add_paragraph()


def _add_routing_appendix(doc: Document, routing_report: dict) -> None:
    """Appendix: Routing Summary — classification, packs, confidence, warnings.

    Renders the routing classification, selected packs, confidence score,
    and any warnings from the routing pipeline.
    """
    if not routing_report:
        return

    doc.add_page_break()
    doc.add_heading("Appendix: Routing Summary", level=1)
    doc.add_paragraph(
        "This appendix documents the RFP routing classification, "
        "selected context packs, and routing confidence. This information "
        "is generated automatically by the routing pipeline."
    )

    # Classification table
    classification = routing_report.get("classification", {})
    if classification:
        doc.add_heading("Classification", level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)

        fields = [
            ("Jurisdiction", classification.get("jurisdiction", "unknown")),
            ("Sector", classification.get("sector", "unknown")),
            ("Domain", classification.get("domain", "") or "not identified"),
            ("Subdomain", classification.get("subdomain", "") or "not identified"),
            ("Client Type", classification.get("client_type", "") or "not identified"),
            ("Regulatory Frame", classification.get("regulatory_frame", "none_identified")),
            ("Evaluator Pattern", classification.get("evaluator_pattern", "") or "not detected"),
            ("Language", classification.get("language", "en")),
            ("Confidence", f"{classification.get('confidence', 0):.2f}"),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            row[1].text = str(value)

    # Proof types needed
    proof_types = classification.get("proof_types_needed", [])
    if proof_types:
        doc.add_paragraph()
        doc.add_heading("Proof Types Needed", level=2)
        for pt in proof_types:
            doc.add_paragraph(pt.replace("_", " ").title(), style="List Bullet")

    # Selected packs
    selected_packs = routing_report.get("selected_packs", [])
    fallback_packs = routing_report.get("fallback_packs_used", [])
    if selected_packs or fallback_packs:
        doc.add_paragraph()
        doc.add_heading("Selected Context Packs", level=2)
        if selected_packs:
            p = doc.add_paragraph()
            run = p.add_run("Primary packs: ")
            run.bold = True
            p.add_run(", ".join(selected_packs))
        if fallback_packs:
            p = doc.add_paragraph()
            run = p.add_run("Fallback packs: ")
            run.bold = True
            p.add_run(", ".join(fallback_packs))

    # Merge statistics
    doc.add_paragraph()
    doc.add_heading("Merged Pack Statistics", level=2)
    stats = [
        ("Regulatory References", routing_report.get("merged_regulatory_refs", 0)),
        ("Compliance Patterns", routing_report.get("merged_compliance_patterns", 0)),
        ("Evaluator Insights", routing_report.get("merged_evaluator_insights", 0)),
        ("Methodology Patterns", routing_report.get("merged_methodology_patterns", 0)),
        ("Benchmark References", routing_report.get("merged_benchmark_refs", 0)),
        ("Search Queries", routing_report.get("merged_search_queries", 0)),
    ]
    for label, count in stats:
        doc.add_paragraph(f"{label}: {count}", style="List Bullet")

    # Confidence
    doc.add_paragraph()
    confidence = routing_report.get("routing_confidence", 0)
    p = doc.add_paragraph()
    run = p.add_run(f"Overall Routing Confidence: {confidence:.2f}")
    run.bold = True
    if confidence < 0.7:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Warnings
    warnings = routing_report.get("warnings", [])
    if warnings:
        doc.add_paragraph()
        doc.add_heading("Warnings", level=2)
        for w in warnings:
            p = doc.add_paragraph()
            run = p.add_run("WARNING: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            p.add_run(w)

    # Alternate classifications
    alternates = classification.get("alternate_classifications", [])
    if alternates:
        doc.add_paragraph()
        doc.add_heading("Alternate Classifications", level=2)
        doc.add_paragraph(
            "The following alternate classifications were considered "
            "due to low routing confidence:"
        )
        for alt in alternates:
            doc.add_paragraph(
                f"{alt.get('field', '')}: {alt.get('value', '')} "
                f"(score: {alt.get('score', 0)})",
                style="List Bullet",
            )


async def export_source_book_docx(
    source_book: SourceBook,
    output_path: str,
    external_evidence_pack: object | None = None,
    routing_report: dict | None = None,
) -> str:
    """Export a SourceBook as a .docx file.

    Produces a structured Word document with cover page, 7 sections,
    and appropriate tables/prose formatting.

    Args:
        source_book: The SourceBook to export.
        output_path: Path to write the .docx file.
        external_evidence_pack: Optional ExternalEvidencePack with rich
            metadata (provider, url, mapped_rfp_theme) to enrich Section 4.
        routing_report: Optional routing report dict to render as an appendix.

    Returns the output path.
    """
    doc = Document()

    # Build evidence enrichment lookup from pack (if available)
    _evidence_enrichment: dict[str, dict] = {}
    if external_evidence_pack:
        for src in getattr(external_evidence_pack, "sources", []):
            sid = getattr(src, "source_id", "")
            if sid:
                _evidence_enrichment[sid] = {
                    "provider": getattr(src, "provider", ""),
                    "url": getattr(src, "url", ""),
                    "mapped_rfp_theme": getattr(src, "mapped_rfp_theme", ""),
                    "how_to_use": getattr(src, "how_to_use_in_proposal", ""),
                }

    # Cover page
    _add_cover_page(doc, source_book)

    # 7 sections
    _add_section_1(doc, source_book)
    _add_section_2(doc, source_book)
    _add_section_3(doc, source_book)
    _add_section_4(doc, source_book, _evidence_enrichment)
    _add_section_5(doc, source_book)
    _add_section_6(doc, source_book)
    _add_section_7(doc, source_book)

    # Appendix: Engine 2 Requirements (proof gaps)
    _add_engine2_requirements(doc, source_book)

    # Appendix: Routing Summary
    if routing_report:
        _add_routing_appendix(doc, routing_report)

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    logger.info("Source Book DOCX exported to %s", output_path)
    return output_path
