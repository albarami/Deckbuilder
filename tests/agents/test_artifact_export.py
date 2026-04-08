"""Artifact-level tests for Engine 1 DOCX export discipline.

These tests generate actual DOCX files from populated SourceBook objects
and verify the rendered document structure, NOT just module/unit helpers.

Acceptance criteria verified:
1. Matrix-first opening for high-density RFPs
2. Structured Likely Evaluation Model table
3. Softened inference language (no inference-as-fact)
4. No unsupported absolutes in final export
5. Typed-field-driven output shape
6. Benchmark framing as analogues, not proof
7. Compliance matrix rendered from typed rows
8. Delivery-control matrix rendered for high density
9. Assumptions & ambiguities table present
10. Sanitized prose passes through classifier
"""

import asyncio
import os
import tempfile

import pytest
from docx import Document

from src.models.source_book import (
    AssertionLabel,
    ClassifiedClaim,
    ComplianceRow,
    DeliveryControlRow,
    EvaluationHypothesis,
    EvidenceLedger,
    EvidenceLedgerEntry,
    ExternalEvidenceSection,
    ProposedSolution,
    RFPInterpretation,
    SourceBook,
    WhyStrategicGears,
)
from src.services.source_book_export import export_source_book_docx


def _build_high_density_source_book() -> SourceBook:
    """Build a realistic high-density source book with all typed fields."""
    return SourceBook(
        client_name="Test Ministry",
        rfp_name="Digital Transformation Advisory",
        language="en",
        requirement_density="high",
        rfp_interpretation=RFPInterpretation(
            objective_and_scope=(
                "The Ministry seeks a consulting partner for enterprise "
                "digital transformation. SG guarantees flawless delivery "
                "with zero learning curve and zero unplanned downtime."
            ),
            constraints_and_compliance=(
                "Bidders must comply with local data sovereignty. "
                "The RFP states that weekly reports are required and "
                "the scoring will favor experience. "
                "The RFP requires the evaluation committee will review "
                "technical submissions first. "
                "Section 3.1 requires team CVs and the evaluators will "
                "prioritize local experience. "
                "According to the RFP, weekly reports are required and "
                "the evaluators will prioritize local firms. "
                "The evaluation committee will almost certainly include "
                "procurement patterns consistently require local content. "
                "The client expects weekly status updates."
            ),
            unstated_evaluator_priorities=(
                "The evaluators will prioritize firms with proven local presence."
            ),
            probable_scoring_logic=(
                "Projected 65/35 scoring logic with technical weighted higher. "
                "The scoring will favor firms with government experience."
            ),
            key_compliance_requirements=[
                "COMP-001 | ISO 27001 certification required",
                "COMP-002 | Minimum 5 TOGAF-certified architects",
            ],
            explicit_requirements=[
                ClassifiedClaim(
                    claim_text="ISO 27001 certification mandatory per Section 4.2",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                    basis="RFP Section 4.2",
                    confidence="high",
                ),
                ClassifiedClaim(
                    claim_text="Minimum 10 years consulting experience",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                    basis="RFP Section 3.1",
                    confidence="high",
                ),
            ],
            inferred_requirements=[
                ClassifiedClaim(
                    claim_text="Saudization ratio likely weighted in evaluation",
                    label=AssertionLabel.INFERENCE,
                    basis="Government procurement patterns",
                    confidence="medium",
                ),
            ],
            external_support=[
                ClassifiedClaim(
                    claim_text="McKinsey Digital Gov report supports phased approach",
                    label=AssertionLabel.EXTERNAL_BENCHMARK,
                    basis="EXT-001 McKinsey 2023",
                    confidence="medium",
                ),
            ],
            assumptions=[
                "Evaluation committee includes technical assessors",
                "Budget envelope sufficient for full scope delivery",
            ],
            ambiguities=[
                "Cloud hosting scope unclear — on-prem vs hybrid not specified",
                "SLA penalty structure not defined in RFP",
            ],
            compliance_rows=[
                ComplianceRow(
                    requirement_id="COMP-001",
                    requirement_text="ISO 27001 certification",
                    sg_response="SG holds ISO 27001:2022",
                    evidence_ref="CLM-0015",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                ),
                ComplianceRow(
                    requirement_id="COMP-002",
                    requirement_text="Minimum 5 TOGAF architects",
                    sg_response="SG recommends 3 TOGAF-certified, 2 pending",
                    evidence_ref="CLM-0021",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                ),
                ComplianceRow(
                    requirement_id="COMP-003",
                    requirement_text="Data sovereignty compliance",
                    sg_response="All data hosted locally per regulation",
                    evidence_ref="CLM-0033",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                ),
                ComplianceRow(
                    requirement_id="COMP-004",
                    requirement_text="Monthly progress reporting",
                    sg_response="PMO dashboard with weekly/monthly cadence",
                    evidence_ref="CLM-0041",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                ),
                ComplianceRow(
                    requirement_id="COMP-005",
                    requirement_text="Knowledge transfer plan",
                    sg_response="Structured KT program with certification",
                    evidence_ref="CLM-0052",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                ),
            ],
            delivery_control_rows=[
                DeliveryControlRow(
                    control_area="Reporting",
                    rfp_requirement="Monthly status reports",
                    proposed_mechanism="PMO dashboard + steering committee",
                    verification_method="Client sign-off per milestone",
                ),
                DeliveryControlRow(
                    control_area="Quality",
                    rfp_requirement="Deliverable QA process",
                    proposed_mechanism="Two-stage review (internal + client)",
                    verification_method="QA checklist sign-off",
                ),
            ],
            evaluation_hypotheses=[
                EvaluationHypothesis(
                    criterion="Technical approach",
                    basis="Common for government digital transformation RFPs",
                    confidence="medium",
                    label=AssertionLabel.INFERENCE,
                    weight_estimate="~40%",
                ),
                EvaluationHypothesis(
                    criterion="Team qualifications",
                    basis="RFP Section 3 requires team CVs",
                    confidence="high",
                    label=AssertionLabel.DIRECT_RFP_FACT,
                    weight_estimate="25%",
                ),
                EvaluationHypothesis(
                    criterion="Past experience",
                    basis="Typical government evaluation pattern",
                    confidence="low",
                    label=AssertionLabel.INFERENCE,
                    weight_estimate="~20%",
                ),
                EvaluationHypothesis(
                    criterion="Financial proposal",
                    basis="Standard two-envelope system",
                    confidence="medium",
                    label=AssertionLabel.INFERENCE,
                    weight_estimate="~15%",
                ),
            ],
        ),
        external_evidence=ExternalEvidenceSection(
            coverage_assessment=(
                "EXT-001 proves that SG can deliver digital transformation. "
                "This benchmark validates SG experience in the government sector."
            ),
        ),
        proposed_solution=ProposedSolution(
            methodology_overview=(
                "SG proposes a phased approach. The system will function "
                "flawlessly from day one with zero learning curve. "
                "This approach operates with zero unplanned downtime "
                "and guarantees 100% compliance."
            ),
            benchmark_references=[
                ClassifiedClaim(
                    claim_text="McKinsey Digital Gov report supports phased delivery",
                    label=AssertionLabel.EXTERNAL_BENCHMARK,
                    basis="EXT-001",
                    confidence="medium",
                ),
            ],
        ),
    )


def _extract_all_text(doc: Document) -> str:
    """Extract all text content from a DOCX document."""
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_heading_order(doc: Document) -> list[str]:
    """Extract headings in order from the document."""
    headings: list[str] = []
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            headings.append(para.text)
    return headings


def _count_tables_before_heading(doc: Document, heading_text: str) -> int:
    """Count how many tables appear before a specific heading."""
    table_count = 0
    heading_found = False
    element_index = 0
    body = doc.element.body

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "tbl":
            if not heading_found:
                table_count += 1
        elif tag == "p":
            for sub in child:
                sub_tag = sub.tag.split("}")[-1] if "}" in sub.tag else sub.tag
                if sub_tag == "pPr":
                    for style_elem in sub:
                        style_tag = style_elem.tag.split("}")[-1] if "}" in style_elem.tag else style_elem.tag
                        if style_tag == "pStyle":
                            val = style_elem.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
                            )
                            if val.startswith("Heading"):
                                # Check text
                                text_parts = []
                                for r in child.iter():
                                    r_tag = r.tag.split("}")[-1] if "}" in r.tag else r.tag
                                    if r_tag == "t" and r.text:
                                        text_parts.append(r.text)
                                full_text = "".join(text_parts)
                                if heading_text in full_text:
                                    heading_found = True
                                    return table_count

    return table_count


# ──────────────────────────────────────────────────────────────
# 1. Matrix-first opening for high-density RFPs
# ──────────────────────────────────────────────────────────────


class TestMatrixFirstRendering:
    """Verify Section 1 opens with matrices before narrative for high density."""

    @pytest.fixture
    def exported_doc(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_artifact.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        return Document(output_path)

    def test_compliance_matrix_before_narrative(self, exported_doc):
        """Compliance matrix table appears before the 1.1 Objective heading."""
        tables_before_objective = _count_tables_before_heading(
            exported_doc, "1.1 Objective"
        )
        assert tables_before_objective >= 1, (
            "Compliance matrix must appear before 1.1 Objective & Scope"
        )

    def test_headings_show_matrix_subsections(self, exported_doc):
        """Section 1 headings include typed-field subsections."""
        headings = _extract_heading_order(exported_doc)
        s1_headings = [h for h in headings if h.startswith("1.") or "Matrix" in h or "Evaluation" in h or "Assumptions" in h or "Compliance" in h or "Delivery" in h]
        heading_text = " ".join(s1_headings)
        assert "Compliance" in heading_text or any("Compliance" in h for h in headings), (
            "Section 1 must include Compliance matrix heading"
        )

    def test_delivery_control_matrix_rendered(self, exported_doc):
        """Delivery-control matrix appears for high-density RFP."""
        all_text = _extract_all_text(exported_doc)
        assert "Delivery-Control Matrix" in all_text or "Control Area" in all_text

    def test_assumptions_ambiguities_table_rendered(self, exported_doc):
        """Assumptions & Ambiguities table appears in Section 1."""
        all_text = _extract_all_text(exported_doc)
        assert "Assumptions" in all_text
        assert "Ambiguity" in all_text or "Ambiguities" in all_text

    def test_evaluation_model_rendered(self, exported_doc):
        """Likely Evaluation Model table appears in Section 1."""
        all_text = _extract_all_text(exported_doc)
        assert "Likely Evaluation Model" in all_text
        assert "Criterion" in all_text
        assert "Confidence" in all_text
        assert "Classification" in all_text

    def test_classified_requirements_tables(self, exported_doc):
        """Explicit and Inferred requirements tables appear."""
        all_text = _extract_all_text(exported_doc)
        assert "Explicit RFP Requirements" in all_text
        assert "Inferred Requirements" in all_text
        assert "RFP Fact" in all_text
        assert "Inferred" in all_text


# ──────────────────────────────────────────────────────────────
# 2. No unsupported absolutes in final export
# ──────────────────────────────────────────────────────────────


class TestAbsoluteSanitizationInExport:
    """Verify unsupported absolutes are removed from the rendered DOCX."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_absolutes.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc).lower()

    def test_no_guarantees(self, exported_text):
        assert "guarantees" not in exported_text
        assert "guarantee" not in exported_text

    def test_no_flawlessly(self, exported_text):
        assert "flawlessly" not in exported_text
        assert "flawless" not in exported_text

    def test_no_zero_learning_curve(self, exported_text):
        assert "zero learning curve" not in exported_text

    def test_no_zero_unplanned_downtime(self, exported_text):
        assert "zero unplanned downtime" not in exported_text

    def test_no_100_percent_compliance(self, exported_text):
        assert "guarantees 100% compliance" not in exported_text


# ──────────────────────────────────────────────────────────────
# 3. Inference language is softened in export
# ──────────────────────────────────────────────────────────────


class TestInferenceSofteningInExport:
    """Verify inference language is softened in the rendered DOCX."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_inference.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc)

    def test_evaluators_will_softened(self, exported_text):
        text_lower = exported_text.lower()
        assert "the evaluators will prioritize" not in text_lower

    def test_scoring_will_favor_softened(self, exported_text):
        text_lower = exported_text.lower()
        assert "the scoring will favor" not in text_lower

    def test_almost_certainly_softened(self, exported_text):
        text_lower = exported_text.lower()
        assert "almost certainly" not in text_lower

    def test_projected_scoring_softened(self, exported_text):
        text_lower = exported_text.lower()
        assert "projected 65/35 scoring" not in text_lower

    def test_likely_appears_as_qualifier(self, exported_text):
        text_lower = exported_text.lower()
        assert "likely" in text_lower, (
            "Inference language should use 'likely' qualifier"
        )


# ──────────────────────────────────────────────────────────────
# 4. Benchmark framing in export
# ──────────────────────────────────────────────────────────────


class TestBenchmarkFramingInExport:
    """Verify benchmarks framed as analogues, not proof of SG capability."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_benchmarks.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc)

    def test_no_proves_sg(self, exported_text):
        assert "proves that SG" not in exported_text

    def test_no_validates_sg_experience(self, exported_text):
        assert "validates SG experience" not in exported_text

    def test_benchmark_references_table_rendered(self, exported_text):
        assert "Benchmark References" in exported_text or "External Benchmark Support" in exported_text


# ──────────────────────────────────────────────────────────────
# 5. Evaluation model structure in export
# ──────────────────────────────────────────────────────────────


class TestEvaluationModelInExport:
    """Verify Likely Evaluation Model is a structured table, not narrative."""

    @pytest.fixture
    def exported_doc(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_eval_model.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        return Document(output_path)

    def test_evaluation_model_is_table(self, exported_doc):
        """The evaluation model must be rendered as a table, not prose."""
        found_table = False
        for table in exported_doc.tables:
            header_cells = [cell.text for cell in table.rows[0].cells]
            if "Criterion" in header_cells and "Confidence" in header_cells:
                found_table = True
                # Verify all hypotheses have Classification column
                assert "Classification" in header_cells
                break
        assert found_table, "Evaluation model must be rendered as a table with Criterion/Confidence/Classification columns"

    def test_evaluation_note_present(self, exported_doc):
        """Italicized inference note must appear before the table."""
        all_text = _extract_all_text(exported_doc).lower()
        assert "inferred from rfp structure" in all_text


# ──────────────────────────────────────────────────────────────
# 6. Full pipeline: typed fields drive the artifact
# ──────────────────────────────────────────────────────────────


class TestTypedFieldsDriveArtifact:
    """Verify the final DOCX is driven by typed fields, not legacy prose."""

    @pytest.fixture
    def exported_doc_and_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_typed.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return doc, _extract_all_text(doc)

    def test_compliance_row_content_in_table(self, exported_doc_and_text):
        """Typed compliance row content appears in a table."""
        _, text = exported_doc_and_text
        assert "COMP-001" in text
        assert "ISO 27001" in text
        assert "CLM-0015" in text

    def test_delivery_control_content_in_table(self, exported_doc_and_text):
        """Typed delivery control content appears in a table."""
        _, text = exported_doc_and_text
        assert "Reporting" in text
        assert "PMO dashboard" in text

    def test_assumption_content_in_table(self, exported_doc_and_text):
        """Assumptions appear in the structured table."""
        _, text = exported_doc_and_text
        assert "technical assessors" in text

    def test_ambiguity_content_in_table(self, exported_doc_and_text):
        """Ambiguities appear in the structured table."""
        _, text = exported_doc_and_text
        assert "Cloud hosting" in text

    def test_evaluation_hypothesis_content(self, exported_doc_and_text):
        """Evaluation hypothesis rows appear with all fields."""
        _, text = exported_doc_and_text
        assert "Technical approach" in text
        assert "~40%" in text

    def test_medium_density_skips_delivery_matrix(self, tmp_path):
        """Medium density RFP should NOT get delivery-control matrix."""
        sb = _build_high_density_source_book()
        sb.requirement_density = "medium"
        output_path = str(tmp_path / "test_medium.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        text = _extract_all_text(doc)
        assert "Delivery-Control Matrix" not in text

    def test_empty_typed_fields_fall_back_to_narrative(self, tmp_path):
        """Without typed fields, Section 1 falls back to narrative order."""
        sb = SourceBook(
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="Simple scope description.",
                key_compliance_requirements=["Basic requirement"],
            ),
        )
        output_path = str(tmp_path / "test_fallback.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        headings = _extract_heading_order(doc)
        # Filter for sub-headings only (1.x), not the section title (1. RFP...)
        s1_sub_headings = [
            h for h in headings
            if h.startswith("1.") and not h.startswith("1. ")
        ]
        assert s1_sub_headings[0] == "1.1 Objective & Scope", (
            "Without typed fields, narrative should come first"
        )


# ──────────────────────────────────────────────────────────────
# FACT-PRESERVATION TESTS (A-E)
# ──────────────────────────────────────────────────────────────


class TestFactPreservation:
    """A. Explicit RFP facts must NOT be softened to 'likely'."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_fact_preservation.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc)

    def test_rfp_states_weekly_reports_preserved(self, exported_text):
        """'The RFP states that weekly reports are required' must stay factual."""
        assert "the rfp states that weekly reports are required" in exported_text.lower()

    def test_rfp_requires_evaluation_committee_preserved(self, exported_text):
        """'The RFP requires the evaluation committee will...' must NOT become 'will likely'."""
        text_lower = exported_text.lower()
        assert "the rfp requires the evaluation committee will review" in text_lower

    def test_no_likely_inserted_into_rfp_fact(self, exported_text):
        """Source-grounded sentences must not contain injected 'likely'."""
        text_lower = exported_text.lower()
        # The fact sentence "The RFP states that weekly reports are required"
        # must appear verbatim without "likely" injected into it.
        # (Other sentences on the same rendered line may contain "likely".)
        assert "the rfp states that weekly reports are required" in text_lower
        assert "the rfp likely states" not in text_lower
        assert "the rfp states that weekly reports are likely" not in text_lower


class TestExplicitScoringFactPreserved:
    """B. Evaluation hypothesis with explicit RFP basis must keep DIRECT_RFP_FACT."""

    @pytest.fixture
    def exported_doc(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_explicit_scoring.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        return Document(output_path)

    def test_explicit_hypothesis_label_preserved(self, exported_doc):
        """Team qualifications hypothesis with 'RFP Section 3 requires' must be RFP Fact."""
        all_text = _extract_all_text(exported_doc)
        # Find the Team qualifications row
        assert "Team qualifications" in all_text
        # The label column for this row must say "RFP Fact", not "Inferred"
        assert "RFP Fact" in all_text

    def test_explicit_hypothesis_not_coerced_to_inference(self, exported_doc):
        """The label must NOT be coerced to Inferred when basis is explicit."""
        for table in exported_doc.tables:
            header_cells = [cell.text for cell in table.rows[0].cells]
            if "Criterion" in header_cells and "Classification" in header_cells:
                class_idx = header_cells.index("Classification")
                crit_idx = header_cells.index("Criterion")
                for row in table.rows[1:]:
                    cells = [cell.text for cell in row.cells]
                    if cells[crit_idx] == "Team qualifications":
                        assert cells[class_idx] == "RFP Fact", (
                            f"Expected 'RFP Fact' but got '{cells[class_idx]}'"
                        )
                        return
        pytest.fail("Team qualifications row not found in evaluation model table")


class TestInferenceStillSoftened:
    """C. Inference language must still be softened when not source-grounded."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_inference_softened.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc).lower()

    def test_evaluators_will_prioritize_softened(self, exported_text):
        """'Evaluators will prioritize' must become 'likely to prioritize'."""
        assert "the evaluators will prioritize" not in exported_text
        assert "likely to prioritize" in exported_text

    def test_scoring_will_favor_softened(self, exported_text):
        """'The scoring will favor' must become 'likely to favor'."""
        assert "the scoring will favor" not in exported_text
        assert "likely to favor" in exported_text

    def test_non_grounded_client_expects_softened(self, exported_text):
        """'The client expects' without RFP citation must be softened."""
        # The fixture has "The client expects weekly status updates."
        # which has NO "The RFP states" prefix, so it must be softened
        assert "the client likely expects" in exported_text


class TestNoStackedHedging:
    """D. No awkward stacked hedging in the exported text."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_stacked_hedges.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc).lower()

    def test_no_likely_with_high_probability(self, exported_text):
        assert "likely with high probability" not in exported_text

    def test_no_probably_likely(self, exported_text):
        assert "probably likely" not in exported_text

    def test_no_typically_likely_expected(self, exported_text):
        assert "typically likely expected" not in exported_text

    def test_no_will_likely_with_high_probability(self, exported_text):
        assert "will likely with high probability" not in exported_text

    def test_clean_single_hedge_used(self, exported_text):
        """The 'will almost certainly include' should become 'will likely include', not stacked."""
        assert "almost certainly" not in exported_text


class TestMatrixFirstRetained:
    """E. Matrix-first Section 1 behavior is retained after fact-preservation fix."""

    @pytest.fixture
    def exported_doc(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_matrix_retained.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        return Document(output_path)

    def test_compliance_matrix_still_first(self, exported_doc):
        """Compliance matrix must still appear before narrative."""
        tables_before = _count_tables_before_heading(exported_doc, "1.1 Objective")
        assert tables_before >= 1

    def test_delivery_control_still_present(self, exported_doc):
        all_text = _extract_all_text(exported_doc)
        assert "Delivery-Control Matrix" in all_text or "Control Area" in all_text

    def test_evaluation_model_still_present(self, exported_doc):
        all_text = _extract_all_text(exported_doc)
        assert "Likely Evaluation Model" in all_text

    def test_assumptions_ambiguities_still_present(self, exported_doc):
        all_text = _extract_all_text(exported_doc)
        assert "Assumptions" in all_text
        assert "Ambiguity" in all_text or "Ambiguities" in all_text


# ──────────────────────────────────────────────────────────────
# CLAUSE-LEVEL FACT PRESERVATION TESTS (A-D)
# ──────────────────────────────────────────────────────────────


class TestMixedFactInferenceSentence:
    """A. Mixed fact + inference sentence: fact clause preserved, inference softened."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_mixed_clause.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc).lower()

    def test_rfp_fact_clause_preserved(self, exported_text):
        """'The RFP states that weekly reports are required' stays factual."""
        assert "the rfp states that weekly reports are required" in exported_text

    def test_adjacent_inference_clause_softened(self, exported_text):
        """'the scoring will favor' in same sentence is softened."""
        assert "the scoring will favor" not in exported_text
        assert "the scoring is likely to favor" in exported_text

    def test_mixed_sentence_has_both(self, exported_text):
        """Both preserved fact and softened inference appear in the output."""
        assert "the rfp states that weekly reports are required" in exported_text
        assert "likely to favor" in exported_text


class TestSectionNumberCitationPreserved:
    """B. Section-number citations (3.1, 4.2) preserved intact."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_section_num.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc).lower()

    def test_section_3_1_intact(self, exported_text):
        """'Section 3.1' must NOT be split into 'Section 3.' and '1'."""
        assert "section 3.1 requires" in exported_text

    def test_section_fact_clause_preserved(self, exported_text):
        """Fact clause 'Section 3.1 requires team CVs' preserved."""
        assert "section 3.1 requires team cvs" in exported_text

    def test_adjacent_inference_softened(self, exported_text):
        """Inference clause after 'and' is softened."""
        assert "the evaluators are likely to prioritize" in exported_text


class TestAccordingToRFPMixedClause:
    """C. According-to-the-RFP mixed clause: fact preserved, inference softened."""

    @pytest.fixture
    def exported_text(self, tmp_path):
        sb = _build_high_density_source_book()
        output_path = str(tmp_path / "test_according.docx")
        asyncio.get_event_loop().run_until_complete(
            export_source_book_docx(sb, output_path)
        )
        doc = Document(output_path)
        return _extract_all_text(doc).lower()

    def test_according_fact_clause_preserved(self, exported_text):
        """'According to the RFP, weekly reports are required' stays factual."""
        assert "according to the rfp, weekly reports are required" in exported_text

    def test_adjacent_evaluators_softened(self, exported_text):
        """'the evaluators will prioritize' after 'and' is softened."""
        assert "the evaluators are likely to prioritize" in exported_text


class TestArabicMixedClause:
    """D. Arabic explicit-source mixed clause: fact preserved, inference softened."""

    def test_arabic_fact_guard_with_inference(self):
        """وفقاً للكراسة clause preserved; inference clause softened."""
        from src.agents.source_book.assertion_classifier import (
            _enforce_inference_rendering,
        )
        text = (
            "وفقاً للكراسة التقارير مطلوبة and "
            "سيعطي المقيّمون أولوية للخبرة المحلية."
        )
        result, count = _enforce_inference_rendering(text)
        # Fact clause preserved
        assert "وفقاً للكراسة التقارير مطلوبة" in result
        # Inference clause softened
        assert "من المرجح أن يعطي المقيّمون أولوية" in result
        assert count >= 1

    def test_arabic_section_numbers_intact(self):
        """Section numbers in Arabic context remain intact."""
        from src.agents.source_book.assertion_classifier import (
            _enforce_inference_rendering,
        )
        text = "حسب المتطلبات في القسم 3.1 التقارير مطلوبة."
        result, _ = _enforce_inference_rendering(text)
        assert "3.1" in result
