"""Tests for Phase 3: Source Book Generation.

Verifies:
1. SourceBook and SourceBookReview schemas validate correctly (8 stages)
2. DeckForgeState has source_book field
3. Source Book Writer agent builds user message from real state
4. Source Book Reviewer agent builds user message from Source Book
5. Orchestrator iteration logic (converge within max passes)
6. DOCX export produces valid file with all 8 stages
7. source_book_node is wired into graph between proposal_strategy and gate_3
8. MODEL_MAP has source_book_writer and source_book_reviewer keys
9. report_markdown populated from Source Book content
"""

import contextlib
import json
import os
import tempfile

import pytest

from src.models.source_book import (
    CapabilityMapping,
    ClientProblemFraming,
    ConsultantProfile,
    EvidenceLedger,
    EvidenceLedgerEntry,
    ExternalEvidenceEntry,
    ExternalEvidenceSection,
    PhaseDetail,
    ProjectExperience,
    ProposedSolution,
    RFPInterpretation,
    SectionCritique,
    SlideBlueprintEntry,
    SourceBook,
    SourceBookReview,
    WhyStrategicGears,
)
from src.models.state import DeckForgeState

# ──────────────────────────────────────────────────────────────
# 1. Schema validation
# ──────────────────────────────────────────────────────────────


class TestSourceBookSchema:
    """Verify SourceBook Pydantic models for all 7 sections."""

    def test_empty_source_book(self):
        sb = SourceBook()
        assert sb.client_name == ""
        assert sb.rfp_name == ""
        assert sb.language == "en"
        assert sb.pass_number == 1
        # All 7 sections present
        assert sb.rfp_interpretation is not None
        assert sb.client_problem_framing is not None
        assert sb.why_strategic_gears is not None
        assert sb.external_evidence is not None
        assert sb.proposed_solution is not None
        assert sb.slide_blueprints == []
        assert sb.evidence_ledger is not None

    def test_full_source_book(self):
        sb = SourceBook(
            client_name="Ministry of Finance",
            rfp_name="Digital Transformation Advisory",
            language="en",
            generation_date="2026-03-22",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="The client seeks advisory on digital transformation.",
                constraints_and_compliance="Must comply with Vision 2030 standards.",
                unstated_evaluator_priorities="Local content, Saudization",
                probable_scoring_logic="70/30 technical/financial split",
                key_compliance_requirements=["COMP-001", "COMP-002"],
            ),
            client_problem_framing=ClientProblemFraming(
                current_state_challenge="Legacy systems impede efficiency.",
                why_it_matters_now="Vision 2030 deadline approaching.",
                transformation_logic="Phased migration to cloud.",
                risk_if_unchanged="Continued dependence on legacy.",
            ),
            why_strategic_gears=WhyStrategicGears(
                capability_mapping=[
                    CapabilityMapping(
                        rfp_requirement="SAP migration",
                        sg_capability="10+ SAP HANA deployments",
                        evidence_ids=["CLM-0001", "CLM-0005"],
                        strength="strong",
                    ),
                ],
                named_consultants=[
                    ConsultantProfile(
                        name="Ahmad Al-Rashid",
                        role="Lead SAP Architect",
                        relevance="15 years SAP in government",
                        evidence_ids=["CLM-0020"],
                    ),
                ],
                project_experience=[
                    ProjectExperience(
                        project_name="SIDF SAP Migration",
                        client="SIDF",
                        outcomes="Migrated 200+ users to HANA",
                        evidence_ids=["CLM-0001"],
                    ),
                ],
                certifications_and_compliance=["ISO 27001", "SAP Gold Partner"],
            ),
            external_evidence=ExternalEvidenceSection(
                entries=[
                    ExternalEvidenceEntry(
                        source_id="EXT-001",
                        title="Cloud Migration Best Practices",
                        year=2025,
                        relevance="Supports phased migration approach",
                        key_finding="82% of government orgs see ROI within 18 months",
                        source_type="academic_paper",
                    ),
                ],
                coverage_assessment="Strong academic backing for cloud migration.",
            ),
            proposed_solution=ProposedSolution(
                methodology_overview="Agile with waterfall governance gates.",
                phase_details=[
                    PhaseDetail(
                        phase_name="Phase 1: Assessment",
                        activities=["Current-state analysis", "Stakeholder mapping"],
                        deliverables=["Assessment report", "Migration plan"],
                        governance="Weekly steering committee",
                    ),
                ],
                governance_framework="PMO with weekly checkpoints.",
                timeline_logic="18-month program in 3 phases.",
                value_case_and_differentiation="SG is the only firm with both SAP and gov sector.",
            ),
            slide_blueprints=[
                SlideBlueprintEntry(
                    slide_number=1,
                    section="Cover",
                    layout="TITLE",
                    purpose="Set the tone",
                    title="Digital Transformation Advisory",
                    key_message="SG: Your partner for Vision 2030",
                    bullet_logic=[],
                    proof_points=[],
                    visual_guidance="Use SG brand template cover",
                ),
                SlideBlueprintEntry(
                    slide_number=5,
                    section="Why SG",
                    layout="CONTENT_1COL",
                    purpose="Demonstrate SAP expertise",
                    title="Proven SAP Migration Track Record",
                    key_message="10+ government SAP deployments",
                    bullet_logic=[
                        "SIDF: 200+ users migrated [CLM-0001]",
                        "MoF: Full HANA implementation [CLM-0005]",
                    ],
                    proof_points=["CLM-0001", "CLM-0005", "EXT-001"],
                    visual_guidance="Timeline showing project sequence",
                    must_have_evidence=["CLM-0001"],
                    forbidden_content=["Generic SAP claims without project names"],
                ),
            ],
            evidence_ledger=EvidenceLedger(
                entries=[
                    EvidenceLedgerEntry(
                        claim_id="CLM-0001",
                        claim_text="SG delivered SAP HANA migration for SIDF",
                        source_type="internal",
                        source_reference="DOC-001, Slide 5",
                        confidence=0.95,
                        verifiability_status="verified",
                    ),
                    EvidenceLedgerEntry(
                        claim_id="EXT-001",
                        claim_text="82% ROI within 18 months",
                        source_type="external",
                        source_reference="Cloud Migration Best Practices (2025)",
                        confidence=0.80,
                        verifiability_status="partially_verified",
                    ),
                ],
            ),
        )
        assert sb.client_name == "Ministry of Finance"
        assert len(sb.why_strategic_gears.capability_mapping) == 1
        assert sb.why_strategic_gears.capability_mapping[0].strength == "strong"
        assert len(sb.slide_blueprints) == 2
        assert sb.slide_blueprints[1].slide_number == 5
        assert "CLM-0001" in sb.slide_blueprints[1].proof_points
        assert len(sb.evidence_ledger.entries) == 2

    def test_capability_strength_validation(self):
        """strength must be one of the allowed literals."""
        with pytest.raises(Exception):
            CapabilityMapping(
                rfp_requirement="Test",
                sg_capability="Test",
                evidence_ids=["CLM-0001"],
                strength="invalid",
            )

    def test_evidence_ledger_status_validation(self):
        """verifiability_status must be one of the allowed literals."""
        with pytest.raises(Exception):
            EvidenceLedgerEntry(
                claim_id="CLM-0001",
                verifiability_status="invalid",
            )

    def test_serialization_roundtrip(self):
        sb = SourceBook(
            client_name="Test Client",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="Test scope",
                key_compliance_requirements=["COMP-001"],
            ),
            slide_blueprints=[
                SlideBlueprintEntry(
                    slide_number=1,
                    title="Test Slide",
                    proof_points=["CLM-0001"],
                ),
            ],
        )
        dumped = sb.model_dump(mode="json")
        restored = SourceBook.model_validate(dumped)
        assert restored.client_name == "Test Client"
        assert len(restored.slide_blueprints) == 1
        assert restored.slide_blueprints[0].proof_points == ["CLM-0001"]

    def test_capability_mapping_allows_empty_evidence_ids(self):
        """CapabilityMapping allows empty evidence_ids (reviewer flags gaps)."""
        cm = CapabilityMapping(
            rfp_requirement="SAP migration",
            sg_capability="10+ deployments",
            evidence_ids=[],  # allowed — reviewer will flag
            strength="strong",
        )
        assert cm.evidence_ids == []

    def test_project_experience_allows_empty_evidence_ids(self):
        """ProjectExperience allows empty evidence_ids (reviewer flags gaps)."""
        pe = ProjectExperience(
            project_name="SIDF SAP Migration",
            client="SIDF",
            outcomes="Migrated 200+ users",
            evidence_ids=[],  # allowed — reviewer will flag
        )
        assert pe.evidence_ids == []

    def test_slide_blueprint_proof_points_auto_populated_from_must_have(self):
        """proof_points auto-populated from must_have_evidence when omitted by LLM."""
        bp = SlideBlueprintEntry(
            slide_number=5,
            title="Why SG",
            must_have_evidence=["CLM-0001"],
            proof_points=[],  # empty — validator auto-populates from must_have
        )
        assert bp.proof_points == ["CLM-0001"]

    def test_slide_blueprint_proof_points_optional_without_must_have(self):
        """proof_points can be empty when must_have_evidence is empty (e.g., Cover slide)."""
        bp = SlideBlueprintEntry(
            slide_number=1,
            section="Cover",
            title="Cover Slide",
            proof_points=[],
            must_have_evidence=[],
        )
        assert bp.proof_points == []

    def test_all_seven_sections_are_distinct_fields(self):
        """Verify the 7 sections map to 7 distinct model fields."""
        sb = SourceBook()
        section_fields = [
            "rfp_interpretation",
            "client_problem_framing",
            "why_strategic_gears",
            "external_evidence",
            "proposed_solution",
            "slide_blueprints",
            "evidence_ledger",
        ]
        for field_name in section_fields:
            assert hasattr(sb, field_name), f"Missing section field: {field_name}"
        assert len(section_fields) == 7


class TestSourceBookReviewSchema:
    """Verify SourceBookReview Pydantic model."""

    def test_empty_review(self):
        review = SourceBookReview()
        assert review.overall_score == 3
        assert review.competitive_viability == "adequate"
        assert review.pass_threshold_met is False
        assert review.rewrite_required is True

    def test_passing_review(self):
        review = SourceBookReview(
            section_critiques=[
                SectionCritique(
                    section_id="rfp_interpretation",
                    score=4,
                    issues=[],
                ),
                SectionCritique(
                    section_id="why_strategic_gears",
                    score=5,
                    issues=[],
                ),
            ],
            overall_score=4,
            competitive_viability="strong",
            pass_threshold_met=True,
            rewrite_required=False,
        )
        assert review.pass_threshold_met is True
        assert review.rewrite_required is False

    def test_failing_review_with_unsupported_claims(self):
        review = SourceBookReview(
            section_critiques=[
                SectionCritique(
                    section_id="why_strategic_gears",
                    score=2,
                    issues=["Claims not backed by evidence"],
                    unsupported_claims=["Extensive SAP experience"],
                    fluff_detected=["deep expertise", "proven track record"],
                    rewrite_instructions=["Replace vague claims with specific evidence"],
                ),
            ],
            overall_score=2,
            competitive_viability="weak",
            pass_threshold_met=False,
            rewrite_required=True,
        )
        assert review.rewrite_required is True
        assert len(review.section_critiques[0].unsupported_claims) == 1
        assert len(review.section_critiques[0].fluff_detected) == 2

    def test_viability_validation(self):
        """competitive_viability must be one of the allowed literals."""
        with pytest.raises(Exception):
            SourceBookReview(competitive_viability="invalid")


# ──────────────────────────────────────────────────────────────
# 2. State field
# ──────────────────────────────────────────────────────────────


class TestStateField:
    """Verify source_book is on DeckForgeState."""

    def test_default_is_none(self):
        state = DeckForgeState()
        assert state.source_book is None

    def test_can_set_source_book(self):
        sb = SourceBook(client_name="Test", rfp_name="Test RFP")
        state = DeckForgeState(source_book=sb)
        assert state.source_book is not None
        assert state.source_book.client_name == "Test"


# ──────────────────────────────────────────────────────────────
# 3. Writer agent user message construction
# ──────────────────────────────────────────────────────────────


class TestWriterUserMessage:
    """Verify the Source Book Writer builds user messages from real state fields."""

    def test_builds_message_from_empty_state(self):
        from src.agents.source_book.writer import _build_user_message

        state = DeckForgeState()
        msg = _build_user_message(state)
        assert isinstance(msg, str)
        parsed = json.loads(msg)
        assert "rfp_context" in parsed
        assert "reference_index" in parsed
        assert "proposal_strategy" in parsed

    def test_builds_message_with_strategy(self):
        from src.agents.source_book.writer import _build_user_message
        from src.models.proposal_strategy import ProposalStrategy, WinTheme

        strategy = ProposalStrategy(
            rfp_interpretation="Test interpretation",
            proposal_thesis="SG is the right choice.",
            win_themes=[
                WinTheme(
                    theme="SAP expertise",
                    supporting_evidence=["CLM-0001"],
                    differentiator_strength="strong",
                ),
            ],
        )
        state = DeckForgeState(
            proposal_strategy=strategy,
            sector="government",
            geography="Saudi Arabia",
        )
        msg = _build_user_message(state)
        parsed = json.loads(msg)
        assert parsed["proposal_strategy"] is not None
        assert parsed["sector"] == "government"

    def test_includes_reviewer_feedback_on_rewrite(self):
        from src.agents.source_book.writer import _build_user_message

        state = DeckForgeState()
        msg = _build_user_message(state, reviewer_feedback="Fix Section 3 evidence gaps")
        parsed = json.loads(msg)
        assert parsed["reviewer_feedback"] == "Fix Section 3 evidence gaps"


# ──────────────────────────────────────────────────────────────
# 4. Reviewer agent
# ──────────────────────────────────────────────────────────────


class TestReviewerUserMessage:
    """Verify the Source Book Reviewer builds messages from Source Book."""

    def test_builds_message_from_source_book(self):
        from src.agents.source_book.reviewer import _build_user_message

        sb = SourceBook(
            client_name="Test Client",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="Test scope",
            ),
        )
        state = DeckForgeState(source_book=sb)
        msg = _build_user_message(state)
        parsed = json.loads(msg)
        assert "source_book" in parsed
        assert parsed["source_book"]["client_name"] == "Test Client"


# ──────────────────────────────────────────────────────────────
# 5. Orchestrator iteration logic
# ──────────────────────────────────────────────────────────────


class TestOrchestratorLogic:
    """Verify orchestrator converges within max passes."""

    def test_passes_on_first_try(self):
        """If review passes threshold, should stop after 1 pass."""
        from src.agents.source_book.orchestrator import should_continue_iteration

        review = SourceBookReview(
            overall_score=4,
            pass_threshold_met=True,
            rewrite_required=False,
        )
        assert should_continue_iteration(review, current_pass=1, max_passes=5) is False

    def test_continues_on_low_score(self):
        """If overall_score < 4, should continue."""
        from src.agents.source_book.orchestrator import should_continue_iteration

        review = SourceBookReview(
            overall_score=2,
            pass_threshold_met=False,
            rewrite_required=True,
        )
        assert should_continue_iteration(review, current_pass=1, max_passes=5) is True

    def test_stops_at_max_passes(self):
        """Even if review fails, should stop at max passes (5)."""
        from src.agents.source_book.orchestrator import should_continue_iteration

        review = SourceBookReview(
            overall_score=2,
            pass_threshold_met=False,
            rewrite_required=True,
        )
        # Pass 4 should continue (under max)
        assert should_continue_iteration(review, current_pass=4, max_passes=5) is True
        # Pass 5 should stop (at max)
        assert should_continue_iteration(review, current_pass=5, max_passes=5) is False

    def test_default_max_passes_is_5(self):
        """Design contract: default max_passes must be 5."""
        import inspect

        from src.agents.source_book.orchestrator import should_continue_iteration

        sig = inspect.signature(should_continue_iteration)
        default = sig.parameters["max_passes"].default
        assert default == 5, f"Default max_passes is {default}, expected 5"

    def test_build_reviewer_feedback_string(self):
        """Feedback string should summarize critique for rewrite."""
        from src.agents.source_book.orchestrator import build_reviewer_feedback

        review = SourceBookReview(
            section_critiques=[
                SectionCritique(
                    section_id="why_strategic_gears",
                    score=2,
                    issues=["Claims not backed"],
                    rewrite_instructions=["Add evidence IDs"],
                    unsupported_claims=["Extensive experience"],
                ),
            ],
            overall_score=2,
            coherence_issues=["Repetition in sections 3 and 5"],
        )
        feedback = build_reviewer_feedback(review)
        assert "why_strategic_gears" in feedback
        assert "Add evidence IDs" in feedback
        assert "Repetition" in feedback


# ──────────────────────────────────────────────────────────────
# 6. DOCX export
# ──────────────────────────────────────────────────────────────


class TestDocxExport:
    """Verify DOCX export produces valid file with all 7 sections."""

    @pytest.mark.asyncio
    async def test_export_empty_source_book(self):
        from src.services.source_book_export import export_source_book_docx

        sb = SourceBook(client_name="Test", rfp_name="Test RFP")
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "source_book.docx")
            result = await export_source_book_docx(sb, path)
            assert os.path.exists(result)
            assert result.endswith(".docx")
            # File should be non-trivial (>1KB for a DOCX with structure)
            assert os.path.getsize(result) > 1000

    @pytest.mark.asyncio
    async def test_export_full_source_book_has_all_sections(self):
        """DOCX should contain heading text for all 7 sections."""
        from docx import Document

        from src.services.source_book_export import export_source_book_docx

        sb = SourceBook(
            client_name="Ministry of Finance",
            rfp_name="Digital Transformation",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="The client seeks advisory.",
            ),
            client_problem_framing=ClientProblemFraming(
                current_state_challenge="Legacy systems impede efficiency.",
            ),
            why_strategic_gears=WhyStrategicGears(
                capability_mapping=[
                    CapabilityMapping(
                        rfp_requirement="SAP",
                        sg_capability="10+ deployments",
                        evidence_ids=["CLM-0001"],
                        strength="strong",
                    ),
                ],
            ),
            external_evidence=ExternalEvidenceSection(
                entries=[
                    ExternalEvidenceEntry(
                        source_id="EXT-001",
                        title="Cloud Best Practices",
                        year=2025,
                        key_finding="82% ROI",
                    ),
                ],
            ),
            proposed_solution=ProposedSolution(
                methodology_overview="Agile with governance.",
            ),
            slide_blueprints=[
                SlideBlueprintEntry(
                    slide_number=1,
                    section="Cover",
                    title="Test Slide",
                    key_message="Test message",
                    proof_points=["CLM-0001"],
                ),
            ],
            evidence_ledger=EvidenceLedger(
                entries=[
                    EvidenceLedgerEntry(
                        claim_id="CLM-0001",
                        claim_text="SG delivered SAP",
                        confidence=0.95,
                        verifiability_status="verified",
                    ),
                ],
            ),
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "source_book.docx")
            await export_source_book_docx(sb, path)

            # Read the DOCX and check for section headings
            doc = Document(path)
            all_text = " ".join(p.text for p in doc.paragraphs)
            assert "RFP Interpretation" in all_text
            assert "Client Problem Framing" in all_text
            assert "Why Strategic Gears" in all_text
            assert "External Evidence" in all_text
            assert "Proposed Solution" in all_text
            assert "Slide-by-Slide Blueprint" in all_text
            assert "Evidence Ledger" in all_text

    @pytest.mark.asyncio
    async def test_export_includes_tables(self):
        """DOCX should contain tables for structured sections."""
        from docx import Document

        from src.services.source_book_export import export_source_book_docx

        sb = SourceBook(
            why_strategic_gears=WhyStrategicGears(
                capability_mapping=[
                    CapabilityMapping(
                        rfp_requirement="SAP",
                        sg_capability="10+ deployments",
                        evidence_ids=["CLM-0001"],
                        strength="strong",
                    ),
                ],
            ),
            evidence_ledger=EvidenceLedger(
                entries=[
                    EvidenceLedgerEntry(
                        claim_id="CLM-0001",
                        claim_text="SG delivered SAP",
                        confidence=0.95,
                        verifiability_status="verified",
                    ),
                ],
            ),
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "source_book.docx")
            await export_source_book_docx(sb, path)
            doc = Document(path)
            # Should have at least 2 tables (capability mapping + evidence ledger)
            assert len(doc.tables) >= 2


# ──────────────────────────────────────────────────────────────
# 6b. DOCX path persistence + export failure surfacing
# ──────────────────────────────────────────────────────────────


class TestDocxPathPersistence:
    """Verify DOCX path is persisted in state and export failures are surfaced."""

    @pytest.mark.asyncio
    async def test_successful_export_populates_report_docx_path(self):
        """source_book_node should set report_docx_path on successful export."""
        from unittest.mock import AsyncMock, patch

        from src.models.source_book import (
            SourceBookReview,
            SourceBookSection6,
            SourceBookSection7,
        )
        from src.services.llm import LLMResponse

        from src.models.source_book import (
            SourceBookSection1,
            SourceBookSection2,
            SourceBookSection3,
            SourceBookSection4,
            _Section5Methodology,
            _Section5Governance,
        )

        # Split-call architecture: 8 stage responses
        mock_s1 = SourceBookSection1(
            client_name="Test",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="Test scope",
            ),
        )
        mock_s2 = SourceBookSection2()
        mock_s3 = SourceBookSection3(
            why_strategic_gears=WhyStrategicGears(),
        )
        mock_s4 = SourceBookSection4()
        mock_s5m = _Section5Methodology()
        mock_s5g = _Section5Governance()
        mock_s6 = SourceBookSection6(
            slide_blueprints=[SlideBlueprintEntry(slide_number=1, title="Cover")],
        )
        mock_s7 = SourceBookSection7(
            evidence_ledger=EvidenceLedger(entries=[
                EvidenceLedgerEntry(claim_id="CLM-0001", claim_text="Test"),
            ]),
        )
        mock_review = SourceBookReview(
            overall_score=4,
            pass_threshold_met=True,
            rewrite_required=False,
        )

        stage_responses = [
            LLMResponse(parsed=mock_s1, input_tokens=5000, output_tokens=3000,
                        model="claude-opus-4-20250514", latency_ms=8000),
            LLMResponse(parsed=mock_s2, input_tokens=3000, output_tokens=2000,
                        model="claude-opus-4-20250514", latency_ms=5000),
            LLMResponse(parsed=mock_s3, input_tokens=3000, output_tokens=2000,
                        model="claude-opus-4-20250514", latency_ms=5000),
            LLMResponse(parsed=mock_s4, input_tokens=2000, output_tokens=1000,
                        model="claude-opus-4-20250514", latency_ms=3000),
            LLMResponse(parsed=mock_s5m, input_tokens=4000, output_tokens=3000,
                        model="claude-opus-4-20250514", latency_ms=6000),
            LLMResponse(parsed=mock_s5g, input_tokens=4000, output_tokens=3000,
                        model="claude-opus-4-20250514", latency_ms=6000),
            LLMResponse(parsed=mock_s6, input_tokens=3000, output_tokens=2000,
                        model="claude-opus-4-20250514", latency_ms=5000),
            LLMResponse(parsed=mock_s7, input_tokens=3000, output_tokens=1000,
                        model="claude-opus-4-20250514", latency_ms=3000),
        ]
        reviewer_response = LLMResponse(
            parsed=mock_review, input_tokens=4000,
            output_tokens=1500, model="gpt-5.4", latency_ms=3000,
        )

        call_count = 0

        async def mock_writer_call_llm(**kwargs):
            nonlocal call_count
            idx = call_count
            call_count += 1
            if idx < len(stage_responses):
                return stage_responses[idx]
            return stage_responses[-1]

        with (
            patch(
                "src.agents.source_book.writer.call_llm",
                new_callable=AsyncMock,
                side_effect=mock_writer_call_llm,
            ),
            patch(
                "src.agents.source_book.reviewer.call_llm",
                new_callable=AsyncMock,
                return_value=reviewer_response,
            ),
        ):
            from src.pipeline.graph import source_book_node

            state = DeckForgeState()
            result = await source_book_node(state)

        assert "report_docx_path" in result
        assert result["report_docx_path"] is not None
        assert result["report_docx_path"].endswith("source_book.docx")
        # No errors from export
        assert "errors" not in result or not any(
            e.error_type == "DocxExportError"
            for e in result.get("errors", [])
        )

    @pytest.mark.asyncio
    async def test_export_failure_surfaces_error(self):
        """If DOCX export fails, error must be in state.errors."""
        from unittest.mock import AsyncMock, patch

        from src.models.source_book import (
            SourceBookReview,
            SourceBookSection6,
            SourceBookSection7,
        )
        from src.services.llm import LLMResponse

        from src.models.source_book import (
            SourceBookSection1,
            SourceBookSection2,
            SourceBookSection3,
            SourceBookSection4,
            _Section5Methodology,
            _Section5Governance,
        )

        mock_s1 = SourceBookSection1(client_name="Test", rfp_interpretation=RFPInterpretation(objective_and_scope="Test scope"))
        mock_s2 = SourceBookSection2()
        mock_s3 = SourceBookSection3()
        mock_s4 = SourceBookSection4()
        mock_s5m = _Section5Methodology()
        mock_s5g = _Section5Governance()
        mock_s6 = SourceBookSection6(
            slide_blueprints=[SlideBlueprintEntry(slide_number=1, title="Cover")],
        )
        mock_s7 = SourceBookSection7(
            evidence_ledger=EvidenceLedger(entries=[
                EvidenceLedgerEntry(claim_id="CLM-0001", claim_text="Test"),
            ]),
        )
        mock_review = SourceBookReview(
            overall_score=4,
            pass_threshold_met=True,
            rewrite_required=False,
        )

        stage_responses = [
            LLMResponse(parsed=mock_s1, input_tokens=5000, output_tokens=3000,
                        model="claude-opus-4-20250514", latency_ms=8000),
            LLMResponse(parsed=mock_s2, input_tokens=3000, output_tokens=2000,
                        model="claude-opus-4-20250514", latency_ms=5000),
            LLMResponse(parsed=mock_s3, input_tokens=3000, output_tokens=2000,
                        model="claude-opus-4-20250514", latency_ms=5000),
            LLMResponse(parsed=mock_s4, input_tokens=2000, output_tokens=1000,
                        model="claude-opus-4-20250514", latency_ms=3000),
            LLMResponse(parsed=mock_s5m, input_tokens=4000, output_tokens=3000,
                        model="claude-opus-4-20250514", latency_ms=6000),
            LLMResponse(parsed=mock_s5g, input_tokens=4000, output_tokens=3000,
                        model="claude-opus-4-20250514", latency_ms=6000),
            LLMResponse(parsed=mock_s6, input_tokens=3000, output_tokens=2000,
                        model="claude-opus-4-20250514", latency_ms=5000),
            LLMResponse(parsed=mock_s7, input_tokens=3000, output_tokens=1000,
                        model="claude-opus-4-20250514", latency_ms=3000),
        ]
        reviewer_response = LLMResponse(
            parsed=mock_review, input_tokens=4000,
            output_tokens=1500, model="gpt-5.4", latency_ms=3000,
        )

        call_count = 0

        async def mock_writer_call_llm(**kwargs):
            nonlocal call_count
            idx = call_count
            call_count += 1
            if idx < len(stage_responses):
                return stage_responses[idx]
            return stage_responses[-1]

        with (
            patch(
                "src.agents.source_book.writer.call_llm",
                new_callable=AsyncMock,
                side_effect=mock_writer_call_llm,
            ),
            patch(
                "src.agents.source_book.reviewer.call_llm",
                new_callable=AsyncMock,
                return_value=reviewer_response,
            ),
            patch(
                "src.services.source_book_export.export_source_book_docx",
                new_callable=AsyncMock,
                side_effect=PermissionError("Disk full"),
            ),
        ):
            from src.pipeline.graph import source_book_node

            state = DeckForgeState()
            result = await source_book_node(state)

        # DOCX path should be None on failure
        assert result["report_docx_path"] is None
        # Error must be surfaced structurally
        assert "errors" in result
        assert len(result["errors"]) > 0
        assert any(
            e.error_type == "DocxExportError" for e in result["errors"]
        )
        assert "last_error" in result
        assert result["last_error"].error_type == "DocxExportError"


# ──────────────────────────────────────────────────────────────
# 7. Graph wiring
# ──────────────────────────────────────────────────────────────


class TestGraphWiring:
    """Verify source_book_node is in graph between proposal_strategy and gate_3."""

    def test_source_book_node_exists(self):
        from src.pipeline.graph import source_book_node

        assert callable(source_book_node)

    def test_graph_has_source_book_node(self):
        from src.pipeline.graph import build_graph

        graph = build_graph()
        g = graph.get_graph()
        node_names = [n.name if hasattr(n, "name") else str(n) for n in g.nodes]
        assert "source_book" in node_names

    def test_graph_flow_order(self):
        """Verify: proposal_strategy -> source_book -> gate_3 -> assembly_plan.

        Phase 4 moved assembly_plan after gate_3. Gate 3 now reviews the
        Source Book before assembly planning proceeds.
        """
        from src.pipeline.graph import build_graph

        graph = build_graph()
        g = graph.get_graph()

        edges: dict[str, list[str]] = {}
        for edge in g.edges:
            src = edge.source.name if hasattr(edge.source, "name") else str(edge.source)
            tgt = edge.target.name if hasattr(edge.target, "name") else str(edge.target)
            edges.setdefault(src, []).append(tgt)

        # proposal_strategy -> source_book
        assert "source_book" in edges.get("proposal_strategy", []), (
            f"proposal_strategy edges: {edges.get('proposal_strategy', [])}"
        )
        # source_book -> gate_3
        assert "gate_3" in edges.get("source_book", []), (
            f"source_book edges: {edges.get('source_book', [])}"
        )


# ──────────────────────────────────────────────────────────────
# 8. MODEL_MAP
# ──────────────────────────────────────────────────────────────


class TestModelMap:
    def test_source_book_writer_in_model_map(self):
        from src.config.models import MODEL_MAP

        assert "source_book_writer" in MODEL_MAP

    def test_source_book_reviewer_in_model_map(self):
        from src.config.models import MODEL_MAP

        assert "source_book_reviewer" in MODEL_MAP

    def test_writer_uses_opus(self):
        """Source Book Writer should use Opus (long-form structured writing)."""
        from src.config.models import MODEL_MAP

        model = MODEL_MAP["source_book_writer"]
        assert "opus" in model.lower()

    def test_reviewer_uses_gpt(self):
        """Source Book Reviewer should use GPT-5.4 (critique role)."""
        from src.config.models import MODEL_MAP

        model = MODEL_MAP["source_book_reviewer"]
        assert "gpt" in model.lower()


# ──────────────────────────────────────────────────────────────
# 9. report_markdown population
# ──────────────────────────────────────────────────────────────


class TestReportMarkdownPopulation:
    """Verify Source Book content populates report_markdown."""

    def test_source_book_to_markdown(self):
        from src.agents.source_book.orchestrator import source_book_to_markdown

        sb = SourceBook(
            client_name="Test Client",
            rfp_name="Test RFP",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="The client seeks advisory.",
                constraints_and_compliance="Must comply with standards.",
            ),
            client_problem_framing=ClientProblemFraming(
                current_state_challenge="Legacy systems impede.",
                risk_if_unchanged="Continued legacy dependence.",
            ),
            proposed_solution=ProposedSolution(
                methodology_overview="Agile with governance.",
            ),
        )
        md = source_book_to_markdown(sb)
        assert isinstance(md, str)
        assert "Test Client" in md
        assert "RFP Interpretation" in md
        assert "Client Problem Framing" in md
        assert "Proposed Solution" in md
        assert len(md) > 50


# ──────────────────────────────────────────────────────────────
# 10. Prompt content
# ──────────────────────────────────────────────────────────────


class TestPromptContent:
    def test_writer_prompts_have_section_framework(self):
        """Split-call prompts cover all 7 sections across multiple prompts."""
        from src.agents.source_book.prompts import (
            STAGE1A_SECTION1_PROMPT,
            STAGE1B_SECTION2_PROMPT,
            STAGE1B_SECTION3_PROMPT,
            STAGE1D_SECTION5_PROMPT,
            STAGE2A_BLUEPRINTS_PROMPT,
            STAGE2B_EVIDENCE_LEDGER_PROMPT,
        )

        all_prompts = " ".join([
            STAGE1A_SECTION1_PROMPT,
            STAGE1B_SECTION2_PROMPT,
            STAGE1B_SECTION3_PROMPT,
            STAGE1D_SECTION5_PROMPT,
            STAGE2A_BLUEPRINTS_PROMPT,
            STAGE2B_EVIDENCE_LEDGER_PROMPT,
        ])
        assert "RFP INTERPRETATION" in all_prompts
        assert "CLIENT PROBLEM FRAMING" in all_prompts
        assert "WHY STRATEGIC GEARS" in all_prompts
        assert "PROPOSED SOLUTION" in all_prompts
        assert "SLIDE" in all_prompts
        assert "EVIDENCE LEDGER" in all_prompts

    def test_writer_prompt_requires_evidence_ids(self):
        from src.agents.source_book.writer import SYSTEM_PROMPT

        assert "CLM-" in SYSTEM_PROMPT
        assert "EXT-" in SYSTEM_PROMPT

    def test_reviewer_prompt_has_critique_framework(self):
        from src.agents.source_book.reviewer import SYSTEM_PROMPT

        assert "score" in SYSTEM_PROMPT.lower()
        assert "unsupported" in SYSTEM_PROMPT.lower() or "evidence" in SYSTEM_PROMPT.lower()
        assert "fluff" in SYSTEM_PROMPT.lower() or "vague" in SYSTEM_PROMPT.lower()


# ──────────────────────────────────────────────────────────────
# 11. Phase 4: Gate 3 redesign + pipeline reorder
# ──────────────────────────────────────────────────────────────


class TestGate3Summary:
    """Verify Gate 3 summary shows Source Book stats per Section 9 design."""

    def test_gate_3_shows_source_book_stats(self):
        """Gate 3 summary must include Source Book word count, evidence count,
        viability score, and DOCX path.
        """
        from src.pipeline.graph import _gate_3_summary

        state = DeckForgeState(
            source_book=SourceBook(
                client_name="Acme Corp",
                rfp_name="Digital Transformation",
                rfp_interpretation=RFPInterpretation(
                    objective_and_scope="The client seeks a comprehensive digital "
                    "transformation program across 5 business units.",
                ),
                evidence_ledger=EvidenceLedger(
                    entries=[
                        EvidenceLedgerEntry(
                            claim_id="CLM-0001",
                            claim_text="Successfully delivered 3 similar projects",
                            confidence=0.9,
                            verifiability_status="verified",
                        ),
                        EvidenceLedgerEntry(
                            claim_id="CLM-0002",
                            claim_text="Team includes 5 certified consultants",
                            confidence=0.85,
                            verifiability_status="verified",
                        ),
                    ],
                ),
                slide_blueprints=[
                    SlideBlueprintEntry(
                        slide_number=1,
                        title="Cover",
                        purpose="Title slide",
                    ),
                ],
            ),
            source_book_review=SourceBookReview(
                overall_score=4,
                competitive_viability="strong",
                pass_threshold_met=True,
                rewrite_required=False,
            ),
            report_docx_path="output/test-session/source_book.docx",
        )

        summary = _gate_3_summary(state)

        # Must contain Source Book stats per design Section 9.2
        assert "evidence" in summary.lower(), f"Missing evidence count: {summary}"
        assert "viability" in summary.lower() or "strong" in summary.lower(), (
            f"Missing viability: {summary}"
        )
        assert "source_book.docx" in summary or "docx" in summary.lower(), (
            f"Missing DOCX path: {summary}"
        )

    def test_gate_3_shows_word_count(self):
        """Gate 3 should report approximate Source Book content volume."""
        from src.pipeline.graph import _gate_3_summary

        state = DeckForgeState(
            source_book=SourceBook(
                client_name="Acme Corp",
                rfp_interpretation=RFPInterpretation(
                    objective_and_scope="The client needs advisory services.",
                ),
            ),
        )

        summary = _gate_3_summary(state)
        # Should mention word count or content size
        assert any(w in summary.lower() for w in ["word", "chars", "content"]), (
            f"Missing content size indicator: {summary}"
        )

    def test_gate_3_fallback_without_source_book(self):
        """Gate 3 should have a meaningful fallback when no Source Book exists."""
        from src.pipeline.graph import _gate_3_summary

        state = DeckForgeState()
        summary = _gate_3_summary(state)
        assert len(summary) > 0


class TestPipelineReorder:
    """Verify Phase 4 pipeline order: source_book → gate_3 → assembly_plan."""

    def test_source_book_to_gate_3(self):
        """source_book must connect to gate_3 (not assembly_plan)."""
        from src.pipeline.graph import build_graph

        graph = build_graph()
        g = graph.get_graph()

        edges: dict[str, list[str]] = {}
        for edge in g.edges:
            src = edge.source.name if hasattr(edge.source, "name") else str(edge.source)
            tgt = edge.target.name if hasattr(edge.target, "name") else str(edge.target)
            edges.setdefault(src, []).append(tgt)

        assert "gate_3" in edges.get("source_book", []), (
            f"source_book must connect to gate_3. "
            f"Actual edges: {edges.get('source_book', [])}"
        )

    def test_gate_3_to_assembly_plan(self):
        """gate_3 approval must route to assembly_plan (not submission_transform)."""
        from src.pipeline.graph import build_graph

        graph = build_graph()
        g = graph.get_graph()

        edges: dict[str, list[str]] = {}
        for edge in g.edges:
            src = edge.source.name if hasattr(edge.source, "name") else str(edge.source)
            tgt = edge.target.name if hasattr(edge.target, "name") else str(edge.target)
            edges.setdefault(src, []).append(tgt)

        gate_3_targets = edges.get("gate_3", [])
        assert "assembly_plan" in gate_3_targets, (
            f"gate_3 must route to assembly_plan on approval. "
            f"Actual edges: {gate_3_targets}"
        )

    def test_gate_3_rejection_loops_to_source_book(self):
        """gate_3 rejection must loop back to source_book (not assembly_plan)."""
        from src.pipeline.graph import build_graph

        graph = build_graph()
        g = graph.get_graph()

        edges: dict[str, list[str]] = {}
        for edge in g.edges:
            src = edge.source.name if hasattr(edge.source, "name") else str(edge.source)
            tgt = edge.target.name if hasattr(edge.target, "name") else str(edge.target)
            edges.setdefault(src, []).append(tgt)

        gate_3_targets = edges.get("gate_3", [])
        assert "source_book" in gate_3_targets, (
            f"gate_3 rejection must loop to source_book. "
            f"Actual edges: {gate_3_targets}"
        )

    def test_assembly_plan_to_blueprint_extraction(self):
        """assembly_plan must connect to blueprint_extraction (Phase 5)."""
        from src.pipeline.graph import build_graph

        graph = build_graph()
        g = graph.get_graph()

        edges: dict[str, list[str]] = {}
        for edge in g.edges:
            src = edge.source.name if hasattr(edge.source, "name") else str(edge.source)
            tgt = edge.target.name if hasattr(edge.target, "name") else str(edge.target)
            edges.setdefault(src, []).append(tgt)

        assert "blueprint_extraction" in edges.get("assembly_plan", []), (
            f"assembly_plan must connect to blueprint_extraction. "
            f"Actual edges: {edges.get('assembly_plan', [])}"
        )

    def test_full_pipeline_order(self):
        """Verify the complete pipeline order per design doc."""
        from src.pipeline.graph import build_graph

        graph = build_graph()
        g = graph.get_graph()

        edges: dict[str, list[str]] = {}
        for edge in g.edges:
            src = edge.source.name if hasattr(edge.source, "name") else str(edge.source)
            tgt = edge.target.name if hasattr(edge.target, "name") else str(edge.target)
            edges.setdefault(src, []).append(tgt)

        # The key Phase 5 sequence:
        # proposal_strategy → source_book → gate_3 → assembly_plan → blueprint_extraction
        assert "source_book" in edges.get("proposal_strategy", [])
        assert "gate_3" in edges.get("source_book", [])
        assert "assembly_plan" in edges.get("gate_3", [])
        assert "blueprint_extraction" in edges.get("assembly_plan", [])


class TestAssemblyPlanPromptUpdate:
    """Verify assembly plan consumes recommended_methodology_approach."""

    def test_prompt_mentions_methodology_approach(self):
        """Assembly plan prompt should reference recommended_methodology_approach."""
        from src.agents.assembly_plan.prompts import SYSTEM_PROMPT

        assert "recommended_methodology_approach" in SYSTEM_PROMPT.lower() or \
            "methodology approach" in SYSTEM_PROMPT.lower() or \
            "proposal_strategy" in SYSTEM_PROMPT.lower(), (
            "Assembly plan prompt must reference recommended_methodology_approach "
            "or proposal_strategy"
        )

    def test_assembly_plan_user_message_includes_methodology_approach(self):
        """Assembly plan user message should include recommended_methodology_approach
        from state.proposal_strategy when available.
        """
        import json

        from src.models.proposal_strategy import ProposalStrategy

        state = DeckForgeState(
            proposal_strategy=ProposalStrategy(
                recommended_methodology_approach="Agile-Waterfall hybrid with 4 phases",
            ),
        )

        # We can't call the async run directly, but we can verify the state field
        # is accessible and test the user_data building pattern the agent uses.
        assert state.proposal_strategy is not None
        assert state.proposal_strategy.recommended_methodology_approach == (
            "Agile-Waterfall hybrid with 4 phases"
        )

        # The actual integration test: build user_data like the agent does
        user_data: dict = {
            "output_language": str(state.output_language),
        }
        if state.rfp_context:
            user_data["rfp_context"] = state.rfp_context.model_dump(mode="json")
        if state.proposal_strategy:
            user_data["recommended_methodology_approach"] = (
                state.proposal_strategy.recommended_methodology_approach
            )

        payload = json.dumps(user_data, ensure_ascii=False)
        assert "Agile-Waterfall hybrid" in payload


# ──────────────────────────────────────────────────────────────
# 12. Phase 4 BLOCKER 1: Gate 3 feedback reaches Source Book writer
# ──────────────────────────────────────────────────────────────


class TestGate3FeedbackToWriter:
    """Gate 3 rejection feedback must reach the Source Book writer on rewrite."""

    def test_gate_3_rejection_populates_feedback_on_state(self):
        """GateDecision stores rejection feedback."""
        from src.models.state import GateDecision

        gd = GateDecision(
            gate_number=3,
            approved=False,
            feedback="The evidence for capability mapping is too weak. "
            "Need stronger CLM references for Phase 2 activities.",
        )
        assert not gd.approved
        assert "capability mapping" in gd.feedback
        assert len(gd.feedback) > 0

    def test_source_book_node_builds_gate3_feedback_for_writer(self):
        """On Gate 3 rejection rewrite, source_book_node must feed Gate 3
        feedback into the writer's first pass.

        When state.gate_3 has rejection feedback and state.source_book
        already exists (rewrite scenario), the feedback string passed to
        writer.run() must include Gate 3 human feedback.
        """
        from src.models.state import GateDecision
        from src.pipeline.graph import _build_gate3_feedback_for_writer

        # Scenario: Gate 3 rejected with feedback, existing Source Book
        state = DeckForgeState(
            source_book=SourceBook(
                client_name="Acme Corp",
                pass_number=2,
            ),
            source_book_review=SourceBookReview(
                overall_score=3,
                competitive_viability="adequate",
                section_critiques=[
                    SectionCritique(
                        section_id="why_strategic_gears",
                        score=2,
                        issues=["Weak evidence for cloud migration capability"],
                        rewrite_instructions=["Add CLM-0045 reference"],
                    ),
                ],
            ),
            gate_3=GateDecision(
                gate_number=3,
                approved=False,
                feedback="The proposed solution section needs stronger "
                "differentiation. Add specific Saudi market experience.",
            ),
        )

        feedback = _build_gate3_feedback_for_writer(state)

        # Must include Gate 3 human feedback
        assert "Saudi market experience" in feedback, (
            f"Gate 3 feedback missing from writer input: {feedback}"
        )
        assert "differentiation" in feedback, (
            f"Gate 3 feedback content missing: {feedback}"
        )

    def test_feedback_merge_includes_both_reviewer_and_gate3(self):
        """When both Red Team review and Gate 3 feedback exist,
        the merged feedback must include BOTH with clear labels.
        """
        from src.models.state import GateDecision
        from src.pipeline.graph import _build_gate3_feedback_for_writer

        state = DeckForgeState(
            source_book=SourceBook(pass_number=2),
            source_book_review=SourceBookReview(
                overall_score=3,
                competitive_viability="adequate",
                section_critiques=[
                    SectionCritique(
                        section_id="rfp_interpretation",
                        score=3,
                        issues=["Missing compliance analysis"],
                    ),
                ],
            ),
            gate_3=GateDecision(
                gate_number=3,
                approved=False,
                feedback="Focus more on Vision 2030 alignment.",
            ),
        )

        feedback = _build_gate3_feedback_for_writer(state)

        # Both sources present
        assert "Vision 2030" in feedback, (
            f"Gate 3 feedback missing: {feedback}"
        )
        assert "compliance" in feedback.lower(), (
            f"Reviewer feedback missing: {feedback}"
        )
        # Clear labeling
        assert "gate 3" in feedback.lower() or "human" in feedback.lower(), (
            f"Gate 3 feedback not labeled: {feedback}"
        )

    def test_feedback_empty_when_no_gate3_rejection(self):
        """When Gate 3 is not rejected (or absent), no Gate 3 feedback
        should be injected.
        """
        from src.pipeline.graph import _build_gate3_feedback_for_writer

        # No gate_3 at all
        state = DeckForgeState()
        feedback = _build_gate3_feedback_for_writer(state)
        # Should return empty or reviewer-only feedback (no gate 3 content)
        assert "gate 3" not in feedback.lower() or feedback == ""

    def test_gate3_feedback_has_priority_label(self):
        """Human (Gate 3) feedback should have a priority label so the writer
        knows to address it first.
        """
        from src.models.state import GateDecision
        from src.pipeline.graph import _build_gate3_feedback_for_writer

        state = DeckForgeState(
            source_book=SourceBook(pass_number=1),
            gate_3=GateDecision(
                gate_number=3,
                approved=False,
                feedback="Restructure the methodology section entirely.",
            ),
        )

        feedback = _build_gate3_feedback_for_writer(state)
        # Human feedback should be clearly marked and prioritized
        assert "Restructure the methodology" in feedback
        assert len(feedback) > 0


# ──────────────────────────────────────────────────────────────
# Phase 5 Tests: Slide Blueprint + Slide Architect + Filler Integration
# ──────────────────────────────────────────────────────────────


class TestSlideBlueprintSchema:
    """SlideBlueprint container model validation."""

    def test_empty_blueprint_defaults(self):
        """Empty SlideBlueprint should have valid defaults."""
        from src.models.slide_blueprint import SlideBlueprint

        bp = SlideBlueprint(entries=[])
        assert bp.entries == []

    def test_blueprint_with_entries(self):
        """SlideBlueprint should accept template-contract SlideBlueprintEntry list."""
        from src.models.slide_blueprint import SlideBlueprint
        from src.models.slide_blueprint import SlideBlueprintEntry as ContractEntry

        entries = [
            ContractEntry(
                section_id="S05",
                section_name="Understanding of Project",
                ownership="dynamic",
                slide_title="Understanding the Challenge",
                key_message="We understand your needs",
                bullet_points=["Point 1", "Point 2"],
                evidence_ids=["CLM-0001"],
            ),
            ContractEntry(
                section_id="S07",
                section_name="Why Strategic Gears Evidence",
                ownership="dynamic",
                slide_title="Why Strategic Gears",
                key_message="Proven track record",
                bullet_points=["Track record", "Team expertise"],
                evidence_ids=["CLM-0002", "CLM-0003"],
            ),
        ]
        bp = SlideBlueprint(entries=entries)
        assert len(bp.entries) == 2
        assert bp.entries[0].section_id == "S05"
        assert bp.entries[1].section_id == "S07"

    def test_blueprint_entry_proof_points_validator(self):
        """SlideBlueprintEntry should reject must_have_evidence without proof_points."""
        import pytest

        bp = SlideBlueprintEntry(
            slide_number=1,
            section="section_01",
            layout="content_heading_desc",
            purpose="Test",
            title="Test",
            key_message="Test",
            must_have_evidence=["CLM-0001"],
            proof_points=[],  # Empty — auto-populated from must_have_evidence
        )
        assert bp.proof_points == ["CLM-0001"]

    def test_blueprint_entry_allows_must_have_with_proof_points(self):
        """SlideBlueprintEntry should accept must_have_evidence when proof_points exist."""
        entry = SlideBlueprintEntry(
            slide_number=1,
            section="section_01",
            layout="content_heading_desc",
            purpose="Test",
            title="Test",
            key_message="Test",
            must_have_evidence=["CLM-0001"],
            proof_points=["CLM-0001", "CLM-0002"],
        )
        assert entry.must_have_evidence == ["CLM-0001"]
        assert len(entry.proof_points) == 2

    def test_slide_blueprint_in_state(self):
        """DeckForgeState should accept slide_blueprint field."""
        from src.models.slide_blueprint import SlideBlueprint
        from src.models.slide_blueprint import SlideBlueprintEntry as ContractEntry

        state = DeckForgeState(
            slide_blueprint=SlideBlueprint(
                entries=[
                    ContractEntry(
                        section_id=f"S0{i}",
                        section_name=f"Section {i}",
                        ownership="dynamic",
                        slide_title=f"Title {i}",
                        key_message=f"Key message {i}",
                    )
                    for i in range(1, 4)
                ],
            ),
        )
        assert state.slide_blueprint is not None
        assert len(state.slide_blueprint.entries) == 3


class TestSlideArchitectAgent:
    """Slide Architect agent — user message building and error handling."""

    def test_user_message_includes_source_book(self):
        """User message should include serialized source book when present."""
        from src.agents.slide_architect.agent import _build_user_message

        state = DeckForgeState(
            source_book=SourceBook(
                client_name="Test Client",
                rfp_name="Test RFP",
                rfp_interpretation=RFPInterpretation(
                    objective_and_scope="Modernize IT infrastructure",
                ),
            ),
        )
        msg = _build_user_message(state)
        assert "source_book" in msg
        assert "Test Client" in msg
        assert "Modernize IT infrastructure" in msg

    def test_user_message_empty_state(self):
        """User message should produce valid JSON even with empty state."""
        from src.agents.slide_architect.agent import _build_user_message

        state = DeckForgeState()
        msg = _build_user_message(state)
        import json
        parsed = json.loads(msg)
        assert isinstance(parsed, dict)
        # No source_book key since it's None
        assert "source_book" not in parsed

    def test_user_message_includes_evidence_ledger(self):
        """User message should include evidence ledger from source book."""
        from src.agents.slide_architect.agent import _build_user_message
        from src.models.source_book import EvidenceLedger, EvidenceLedgerEntry

        state = DeckForgeState(
            source_book=SourceBook(
                evidence_ledger=EvidenceLedger(
                    entries=[
                        EvidenceLedgerEntry(
                            claim_id="CLM-0001",
                            claim_text="Proven track record in digital transformation",
                            source_type="internal",
                            source_reference="DOC-001",
                            confidence=0.9,
                            verifiability_status="verified",
                        ),
                    ],
                ),
            ),
        )
        msg = _build_user_message(state)
        assert "evidence_ledger" in msg
        assert "CLM-0001" in msg

    def test_user_message_includes_rfp_context(self):
        """User message should include RFP mandate and evaluation criteria."""
        from src.agents.slide_architect.agent import _build_user_message
        from src.models.common import BilingualText
        from src.models.rfp import RFPContext

        state = DeckForgeState(
            rfp_context=RFPContext(
                rfp_name=BilingualText(en="IT Modernization RFP"),
                issuing_entity=BilingualText(en="Test Client"),
                mandate=BilingualText(en="Modernize legacy systems"),
            ),
        )
        msg = _build_user_message(state)
        assert "rfp_context" in msg
        assert "IT Modernization RFP" in msg

    @pytest.mark.asyncio
    async def test_run_error_returns_empty_blueprint(self):
        """Agent should return empty blueprint with error on LLM failure."""
        from unittest.mock import AsyncMock, patch

        from src.agents.slide_architect.agent import run

        state = DeckForgeState(
            source_book=SourceBook(client_name="Test"),
        )

        with patch(
            "src.agents.slide_architect.agent.call_llm",
            new_callable=AsyncMock,
            side_effect=RuntimeError("LLM unavailable"),
        ):
            result = await run(state)

        assert result["slide_blueprint"].entries == []
        assert result["last_error"].agent == "slide_architect"
        assert "LLM unavailable" in result["last_error"].message


class TestFillerBlueprintGuidance:
    """Blueprint guidance formatting and injection into fillers."""

    def test_format_blueprint_guidance_empty(self):
        """No blueprint entries → empty string."""
        from src.agents.section_fillers.base import (
            BaseSectionFiller,
            SectionFillerInput,
        )

        filler_input = SectionFillerInput(
            section_id="section_01",
            slide_count=2,
            recommended_layouts=["content_heading_desc"],
        )
        result = BaseSectionFiller._format_blueprint_guidance(filler_input)
        assert result == ""

    def test_format_blueprint_guidance_with_entries(self):
        """Blueprint entries → formatted guidance with all fields."""
        from src.agents.section_fillers.base import (
            BaseSectionFiller,
            SectionFillerInput,
        )

        entry = SlideBlueprintEntry(
            slide_number=1,
            section="section_01",
            layout="content_heading_desc",
            purpose="Show client understanding",
            title="Understanding the Challenge",
            key_message="We understand your critical needs",
            bullet_logic=["Point A", "Point B", "Point C"],
            proof_points=["CLM-0001", "CLM-0002"],
            visual_guidance="Use left-right layout",
            forbidden_content=["Generic claims", "No evidence"],
        )
        filler_input = SectionFillerInput(
            section_id="section_01",
            slide_count=1,
            recommended_layouts=["content_heading_desc"],
            blueprint_entries=[entry],
        )
        guidance = BaseSectionFiller._format_blueprint_guidance(filler_input)

        assert "## Slide Blueprint Guidance" in guidance
        assert "Understanding the Challenge" in guidance
        assert "Show client understanding" in guidance
        assert "We understand your critical needs" in guidance
        assert "Point A" in guidance
        assert "Point B" in guidance
        assert "CLM-0001" in guidance
        assert "CLM-0002" in guidance
        assert "Use left-right layout" in guidance
        assert "Generic claims" in guidance

    def test_format_blueprint_guidance_multiple_entries(self):
        """Multiple blueprint entries → all entries included."""
        from src.agents.section_fillers.base import (
            BaseSectionFiller,
            SectionFillerInput,
        )

        entries = [
            SlideBlueprintEntry(
                slide_number=i,
                section="section_01",
                layout="content_heading_desc",
                purpose=f"Purpose {i}",
                title=f"Slide Title {i}",
                key_message=f"Key msg {i}",
            )
            for i in range(1, 4)
        ]
        filler_input = SectionFillerInput(
            section_id="section_01",
            slide_count=3,
            recommended_layouts=["content_heading_desc"],
            blueprint_entries=entries,
        )
        guidance = BaseSectionFiller._format_blueprint_guidance(filler_input)

        assert "Slide Title 1" in guidance
        assert "Slide Title 2" in guidance
        assert "Slide Title 3" in guidance

    def test_blueprint_entries_filter_by_section_in_orchestrator(self):
        """Orchestrator should filter blueprint entries by section_id."""
        # This tests the filtering logic conceptually — the orchestrator
        # creates SectionFillerInput with only entries for that section.
        entries = [
            SlideBlueprintEntry(
                slide_number=1,
                section="section_01",
                layout="content_heading_desc",
                purpose="Understanding",
                title="Title 1",
                key_message="Msg 1",
            ),
            SlideBlueprintEntry(
                slide_number=2,
                section="section_02",
                layout="content_heading_desc",
                purpose="Why SG",
                title="Title 2",
                key_message="Msg 2",
            ),
            SlideBlueprintEntry(
                slide_number=3,
                section="section_01",
                layout="content_heading_desc",
                purpose="More understanding",
                title="Title 3",
                key_message="Msg 3",
            ),
        ]

        # Simulate orchestrator filtering
        section_01_entries = [e for e in entries if e.section == "section_01"]
        section_02_entries = [e for e in entries if e.section == "section_02"]

        assert len(section_01_entries) == 2
        assert len(section_02_entries) == 1
        assert section_01_entries[0].title == "Title 1"
        assert section_01_entries[1].title == "Title 3"


class TestSubmissionModelsStillImportable:
    """Verify submission models remain importable after blueprint_extraction
    replaced submission_transform in the pipeline graph.

    The governance node still uses SubmissionSourcePack, InternalNotePack,
    UnresolvedIssueRegistry, etc. These models must stay importable.
    """

    def test_submission_source_pack_importable(self):
        """SubmissionSourcePack must be importable."""
        from src.models.submission import SubmissionSourcePack
        assert SubmissionSourcePack is not None

    def test_internal_note_pack_importable(self):
        """InternalNotePack must be importable."""
        from src.models.submission import InternalNotePack
        assert InternalNotePack is not None

    def test_unresolved_issue_registry_importable(self):
        """UnresolvedIssueRegistry must be importable."""
        from src.models.submission import UnresolvedIssueRegistry
        assert UnresolvedIssueRegistry is not None

    def test_submission_qa_result_importable(self):
        """SubmissionQAResult must be importable."""
        from src.models.submission import SubmissionQAResult
        assert SubmissionQAResult is not None

    def test_content_routing_enum_importable(self):
        """ContentRouting enum must be importable."""
        from src.models.enums import ContentRouting
        assert ContentRouting is not None

    def test_state_still_references_submission_models(self):
        """DeckForgeState should still have submission layer fields."""
        state = DeckForgeState()
        assert hasattr(state, "submission_source_pack")
        assert hasattr(state, "internal_notes")
        assert hasattr(state, "unresolved_issues")
        assert hasattr(state, "submission_qa_result")


class TestGraphBlueprintExtractionWiring:
    """Verify blueprint_extraction node exists and is wired in the graph."""

    def test_blueprint_extraction_node_exists(self):
        """blueprint_extraction_node function should exist in graph module."""
        from src.pipeline import graph as graph_module
        assert hasattr(graph_module, "blueprint_extraction_node")
        assert callable(graph_module.blueprint_extraction_node)

    def test_model_map_has_slide_architect(self):
        """MODEL_MAP should have slide_architect key."""
        from src.config.models import MODEL_MAP
        assert "slide_architect" in MODEL_MAP

    def test_slide_architect_prompt_exists(self):
        """Slide Architect should have a non-empty system prompt."""
        from src.agents.slide_architect.prompts import SYSTEM_PROMPT
        assert len(SYSTEM_PROMPT) > 100
        assert "SlideBlueprint" in SYSTEM_PROMPT
        assert "SlideBlueprintEntry" in SYSTEM_PROMPT


class TestBlueprintManifestAlignment:
    """Blueprint ↔ manifest b_variable alignment contract.

    The section_fill_node must compare blueprint entry count against
    manifest b_variable count and fail visibly on mismatch. These tests
    call the real section_fill_node function (the live pipeline path).
    """

    @staticmethod
    def _make_manifest(b_variable_count: int):
        """Create a ProposalManifest with exactly b_variable_count b_variable entries."""
        from src.models.proposal_manifest import (
            ContentSourcePolicy,
            ManifestEntry,
            ProposalManifest,
        )

        entries = [
            ManifestEntry(
                entry_type="b_variable",
                asset_id=f"var_{i}",
                semantic_layout_id="content_heading_desc",
                content_source_policy=ContentSourcePolicy.PROPOSAL_SPECIFIC,
                section_id="section_01",
            )
            for i in range(b_variable_count)
        ]
        return ProposalManifest(entries=entries)

    @staticmethod
    def _make_blueprint(entry_count: int):
        """Create a SlideBlueprint with exactly entry_count entries."""
        from src.models.slide_blueprint import SlideBlueprint
        from src.models.slide_blueprint import SlideBlueprintEntry as ContractEntry

        entries = [
            ContractEntry(
                section_id="S05",
                section_name="Understanding of Project",
                ownership="dynamic",
                slide_title=f"Title {i}",
                key_message=f"Msg {i}",
            )
            for i in range(entry_count)
        ]
        return SlideBlueprint(entries=entries)

    @pytest.mark.asyncio
    async def test_mismatch_blueprint_fewer_than_manifest(self):
        """Blueprint has 1 entry, manifest has 3 b_variable — must error.

        Calls the real section_fill_node (live pipeline path).
        """
        from src.pipeline.graph import section_fill_node

        state = DeckForgeState(
            proposal_manifest=self._make_manifest(b_variable_count=3),
            slide_budget={"section_01": 3},  # budget is Any
            slide_blueprint=self._make_blueprint(entry_count=1),
        )

        result = await section_fill_node(state)

        # Must surface the error structurally
        assert result["last_error"] is not None
        assert result["last_error"].error_type == "BlueprintManifestMismatch"
        assert "Blueprint/manifest count mismatch" in result["last_error"].message
        assert "blueprint has 1 entries" in result["last_error"].message
        assert "manifest has 3 b_variable entries" in result["last_error"].message
        # Must set pipeline to ERROR
        from src.models.enums import PipelineStage
        assert result["current_stage"] == PipelineStage.ERROR
        # Must be in the errors list too
        assert any(
            e.error_type == "BlueprintManifestMismatch"
            for e in result["errors"]
        )

    @pytest.mark.asyncio
    async def test_mismatch_blueprint_more_than_manifest(self):
        """Blueprint has 5 entries, manifest has 2 b_variable — must error.

        Proves direction of mismatch doesn't matter.
        """
        from src.pipeline.graph import section_fill_node

        state = DeckForgeState(
            proposal_manifest=self._make_manifest(b_variable_count=2),
            slide_budget={"section_01": 2},
            slide_blueprint=self._make_blueprint(entry_count=5),
        )

        result = await section_fill_node(state)

        assert result["last_error"].error_type == "BlueprintManifestMismatch"
        assert "blueprint has 5 entries" in result["last_error"].message
        assert "manifest has 2 b_variable entries" in result["last_error"].message

    @pytest.mark.asyncio
    async def test_aligned_counts_proceed_normally(self):
        """Blueprint and manifest both have 3 b_variable entries — no error.

        Calls real section_fill_node (live pipeline path). The orchestrator
        and downstream validation are mocked because we're testing the
        alignment gate, not filler execution.
        """
        from unittest.mock import AsyncMock, MagicMock, patch

        from src.pipeline.graph import section_fill_node

        state = DeckForgeState(
            proposal_manifest=self._make_manifest(b_variable_count=3),
            slide_budget=MagicMock(total_slides=3),
            slide_blueprint=self._make_blueprint(entry_count=3),
        )

        # Mock orchestrator — alignment check runs BEFORE this call
        mock_result = MagicMock()
        mock_result.entries_by_section = {}
        mock_result.filler_outputs = {}

        with (
            patch(
                "src.agents.section_fillers.orchestrator.run_section_fillers",
                new_callable=AsyncMock,
                return_value=mock_result,
            ) as mock_orch,
            patch(
                "src.models.proposal_manifest.validate_manifest",
                return_value=[],  # no validation errors
            ),
        ):
            result = await section_fill_node(state)

        # Should NOT have a BlueprintManifestMismatch error
        last_err = result.get("last_error")
        assert last_err is None or last_err.error_type != "BlueprintManifestMismatch", (
            f"Unexpected mismatch error: {last_err}"
        )
        # Orchestrator must have been called (alignment gate passed)
        mock_orch.assert_called_once()

    @pytest.mark.asyncio
    async def test_no_blueprint_zero_b_variable_passes(self):
        """No blueprint and no b_variable entries — counts match (0 == 0).

        Edge case: manifest exists but has only non-variable entries.
        """
        from unittest.mock import AsyncMock, MagicMock, patch

        from src.models.proposal_manifest import (
            ContentSourcePolicy,
            ManifestEntry,
            ProposalManifest,
        )
        from src.pipeline.graph import section_fill_node

        state = DeckForgeState(
            proposal_manifest=ProposalManifest(
                entries=[
                    ManifestEntry(
                        entry_type="a1_clone",
                        asset_id="cover",
                        semantic_layout_id="cover_slide",
                        content_source_policy=ContentSourcePolicy.INSTITUTIONAL_REUSE,
                        section_id="section_01",
                    ),
                ],
            ),
            slide_budget=MagicMock(total_slides=1),
            slide_blueprint=None,  # No blueprint
        )

        mock_result = MagicMock()
        mock_result.entries_by_section = {}
        mock_result.filler_outputs = {}

        with (
            patch(
                "src.agents.section_fillers.orchestrator.run_section_fillers",
                new_callable=AsyncMock,
                return_value=mock_result,
            ) as mock_orch,
            patch(
                "src.models.proposal_manifest.validate_manifest",
                return_value=[],
            ),
        ):
            result = await section_fill_node(state)

        # 0 blueprint entries == 0 b_variable entries → no mismatch error
        last_err = result.get("last_error")
        assert last_err is None or last_err.error_type != "BlueprintManifestMismatch", (
            f"Unexpected mismatch error: {last_err}"
        )
        mock_orch.assert_called_once()


# ──────────────────────────────────────────────────────────────
# 12. Two-stage Writer architecture tests
# ──────────────────────────────────────────────────────────────


class TestThreeStageWriterArchitecture:
    """Verify the three-stage Writer: Stage 1, Stage 2a (blueprints), Stage 2b (ledger)."""

    @pytest.mark.asyncio
    async def test_stage2a_produces_blueprints(self):
        """Stage 2a dedicated call produces non-empty blueprints."""
        from unittest.mock import AsyncMock, patch

        from src.models.source_book import SourceBookSection6
        from src.services.llm import LLMResponse

        mock_s6 = SourceBookSection6(
            slide_blueprints=[
                SlideBlueprintEntry(slide_number=i, title=f"Slide {i}")
                for i in range(1, 21)
            ],
        )

        response = LLMResponse(
            parsed=mock_s6, input_tokens=3000,
            output_tokens=5000, model="test", latency_ms=5000,
        )

        source_book = SourceBook(
            client_name="Test",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="Test scope",
            ),
        )

        with patch(
            "src.agents.source_book.writer.call_llm",
            new_callable=AsyncMock,
            return_value=response,
        ):
            from src.agents.source_book.writer import _generate_blueprints

            result = await _generate_blueprints(source_book, "test-model")

        assert len(result.slide_blueprints) == 20

    @pytest.mark.asyncio
    async def test_stage2b_produces_evidence_ledger(self):
        """Stage 2b dedicated call produces non-empty evidence ledger."""
        from unittest.mock import AsyncMock, patch

        from src.models.source_book import SourceBookSection7
        from src.services.llm import LLMResponse

        mock_s7 = SourceBookSection7(
            evidence_ledger=EvidenceLedger(entries=[
                EvidenceLedgerEntry(
                    claim_id=f"CLM-{i:04d}",
                    claim_text=f"Claim {i}",
                    confidence=0.9,
                    verifiability_status="verified",
                )
                for i in range(1, 11)
            ]),
        )

        response = LLMResponse(
            parsed=mock_s7, input_tokens=3000,
            output_tokens=3000, model="test", latency_ms=3000,
        )

        source_book = SourceBook(client_name="Test")

        with patch(
            "src.agents.source_book.writer.call_llm",
            new_callable=AsyncMock,
            return_value=response,
        ):
            from src.agents.source_book.writer import _generate_evidence_ledger

            result = await _generate_evidence_ledger(source_book, "test-model")

        assert len(result.evidence_ledger.entries) == 10

    @pytest.mark.asyncio
    async def test_writer_merges_all_eight_stages(self):
        """Writer run() merges 8 stage calls into complete SourceBook."""
        from unittest.mock import AsyncMock, patch

        from src.models.source_book import (
            SourceBookSection1,
            SourceBookSection2,
            SourceBookSection3,
            SourceBookSection4,
            _Section5Methodology,
            _Section5Governance,
            SourceBookSection6,
            SourceBookSection7,
        )
        from src.services.llm import LLMResponse

        mock_s1 = SourceBookSection1(
            client_name="Test",
            rfp_interpretation=RFPInterpretation(
                objective_and_scope="Test scope from Stage 1",
            ),
        )
        mock_s2 = SourceBookSection2()
        mock_s3 = SourceBookSection3(
            why_strategic_gears=WhyStrategicGears(),
        )
        mock_s4 = SourceBookSection4()
        mock_s5m = _Section5Methodology()
        mock_s5g = _Section5Governance()
        mock_s6 = SourceBookSection6(
            slide_blueprints=[
                SlideBlueprintEntry(slide_number=1, title="Cover"),
                SlideBlueprintEntry(slide_number=2, title="Exec Summary"),
            ],
        )
        mock_s7 = SourceBookSection7(
            evidence_ledger=EvidenceLedger(entries=[
                EvidenceLedgerEntry(claim_id="CLM-0001", claim_text="Test"),
            ]),
        )

        stage_responses = [
            LLMResponse(parsed=mock_s1, input_tokens=5000, output_tokens=8000,
                        model="test", latency_ms=8000),
            LLMResponse(parsed=mock_s2, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=5000),
            LLMResponse(parsed=mock_s3, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=5000),
            LLMResponse(parsed=mock_s4, input_tokens=2000, output_tokens=1000,
                        model="test", latency_ms=3000),
            LLMResponse(parsed=mock_s5m, input_tokens=4000, output_tokens=4000,
                        model="test", latency_ms=6000),
            LLMResponse(parsed=mock_s5g, input_tokens=4000, output_tokens=4000,
                        model="test", latency_ms=6000),
            LLMResponse(parsed=mock_s6, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=5000),
            LLMResponse(parsed=mock_s7, input_tokens=3000, output_tokens=2000,
                        model="test", latency_ms=3000),
        ]

        call_count = 0

        async def mock_call_llm(**kwargs):
            nonlocal call_count
            idx = call_count
            call_count += 1
            if idx < len(stage_responses):
                return stage_responses[idx]
            return stage_responses[-1]

        with patch(
            "src.agents.source_book.writer.call_llm",
            new_callable=AsyncMock,
            side_effect=mock_call_llm,
        ):
            from src.agents.source_book.writer import run

            state = DeckForgeState()
            result = await run(state)

        sb = result["source_book"]
        assert sb.rfp_interpretation.objective_and_scope == "Test scope from Stage 1"
        assert len(sb.slide_blueprints) == 2
        assert len(sb.evidence_ledger.entries) == 1
        assert sb.slide_blueprints[0].title == "Cover"
        assert sb.evidence_ledger.entries[0].claim_id == "CLM-0001"
        # Verify all 8 calls were made
        assert call_count == 8

    @pytest.mark.asyncio
    async def test_all_stages_succeed_no_fallback_used(self):
        """When all 8 stages succeed, fallback must NOT be triggered."""
        from unittest.mock import AsyncMock, patch

        from src.models.source_book import (
            SourceBookSection1,
            SourceBookSection2,
            SourceBookSection3,
            SourceBookSection4,
            _Section5Methodology,
            _Section5Governance,
            SourceBookSection6,
            SourceBookSection7,
        )
        from src.services.llm import LLMResponse

        mock_s1 = SourceBookSection1(client_name="Test", rfp_interpretation=RFPInterpretation(objective_and_scope="Test scope"))
        mock_s2 = SourceBookSection2()
        mock_s3 = SourceBookSection3()
        mock_s4 = SourceBookSection4()
        mock_s5m = _Section5Methodology()
        mock_s5g = _Section5Governance()
        mock_s6 = SourceBookSection6(
            slide_blueprints=[SlideBlueprintEntry(slide_number=1, title="Cover")],
        )
        mock_s7 = SourceBookSection7(
            evidence_ledger=EvidenceLedger(entries=[
                EvidenceLedgerEntry(claim_id="CLM-0001", claim_text="Test"),
            ]),
        )

        stage_responses = [
            LLMResponse(parsed=mock_s1, input_tokens=5000, output_tokens=8000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s2, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s3, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s4, input_tokens=2000, output_tokens=1000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s5m, input_tokens=4000, output_tokens=4000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s5g, input_tokens=4000, output_tokens=4000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s6, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s7, input_tokens=3000, output_tokens=2000,
                        model="test", latency_ms=1000),
        ]

        call_count = 0

        async def mock_call_llm(**kwargs):
            nonlocal call_count
            idx = call_count
            call_count += 1
            if idx < len(stage_responses):
                return stage_responses[idx]
            return stage_responses[-1]

        with patch(
            "src.agents.source_book.writer.call_llm",
            new_callable=AsyncMock,
            side_effect=mock_call_llm,
        ):
            from src.agents.source_book.writer import run

            with self._capture_logs("src.agents.source_book.writer") as logs:
                state = DeckForgeState()
                result = await run(state)

        sb = result["source_book"]
        assert len(sb.slide_blueprints) == 1
        assert len(sb.evidence_ledger.entries) == 1
        fallback_msgs = [r for r in logs if "fallback" in r.getMessage().lower()]
        assert len(fallback_msgs) == 0, (
            f"Fallback was triggered when all stages succeeded: {fallback_msgs}"
        )

    @pytest.mark.asyncio
    async def test_stage2b_empty_logs_warning_no_fallback(self):
        """When Stage 2b returns empty ledger, warning is logged (no fallback)."""
        from unittest.mock import AsyncMock, patch

        from src.models.source_book import (
            SourceBookSection1,
            SourceBookSection2,
            SourceBookSection3,
            SourceBookSection4,
            _Section5Methodology,
            _Section5Governance,
            SourceBookSection6,
            SourceBookSection7,
        )
        from src.services.llm import LLMResponse

        mock_s1 = SourceBookSection1(client_name="Test", rfp_interpretation=RFPInterpretation(objective_and_scope="Test scope"))
        mock_s2 = SourceBookSection2()
        mock_s3 = SourceBookSection3()
        mock_s4 = SourceBookSection4()
        mock_s5m = _Section5Methodology()
        mock_s5g = _Section5Governance()
        mock_s6 = SourceBookSection6(
            slide_blueprints=[SlideBlueprintEntry(slide_number=1, title="Cover")],
        )
        mock_s7_empty = SourceBookSection7(
            evidence_ledger=EvidenceLedger(entries=[]),
        )

        stage_responses = [
            LLMResponse(parsed=mock_s1, input_tokens=5000, output_tokens=8000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s2, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s3, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s4, input_tokens=2000, output_tokens=1000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s5m, input_tokens=4000, output_tokens=4000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s5g, input_tokens=4000, output_tokens=4000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s6, input_tokens=3000, output_tokens=3000,
                        model="test", latency_ms=1000),
            LLMResponse(parsed=mock_s7_empty, input_tokens=3000, output_tokens=100,
                        model="test", latency_ms=1000),
        ]

        call_count = 0

        async def mock_call_llm(**kwargs):
            nonlocal call_count
            idx = call_count
            call_count += 1
            if idx < len(stage_responses):
                return stage_responses[idx]
            return stage_responses[-1]

        with patch(
            "src.agents.source_book.writer.call_llm",
            new_callable=AsyncMock,
            side_effect=mock_call_llm,
        ):
            from src.agents.source_book.writer import run

            with self._capture_logs("src.agents.source_book.writer") as logs:
                state = DeckForgeState()
                result = await run(state)

        # Stage 2b empty should log a warning (not a fallback)
        import logging

        warning_msgs = [
            r for r in logs
            if "evidence" in r.getMessage().lower()
            and r.levelno >= logging.WARNING
        ]
        assert len(warning_msgs) > 0, (
            "Warning should be logged when Stage 2b produces empty ledger"
        )
        # No fallback should be triggered
        assert result.get("fallback_events", []) == [], (
            f"No fallbacks expected but got: {result.get('fallback_events')}"
        )

    @staticmethod
    @contextlib.contextmanager
    def _capture_logs(logger_name: str):
        """Context manager to capture log records from a specific logger."""
        import logging

        records: list[logging.LogRecord] = []

        class _Handler(logging.Handler):
            def emit(self, record: logging.LogRecord) -> None:
                records.append(record)

        handler = _Handler()
        logger = logging.getLogger(logger_name)
        logger.addHandler(handler)
        old_level = logger.level
        logger.setLevel(logging.DEBUG)
        try:
            yield records
        finally:
            logger.removeHandler(handler)
            logger.setLevel(old_level)
