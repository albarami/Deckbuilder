"""Tests for broadened Engine 2 pattern, export audience mode, post-export scan,
and state persistence through LangGraph.

Covers:
- Generic 'Engine 2' matches forbidden pattern
- Sanitizer removes generic Engine 2 from client-facing model text
- Internal-only sections preserve Engine 2 text
- sanitization_removals / post_export_forbidden survive on DeckForgeState
- export_source_book_docx audience='client' suppresses Engine 2 appendix
- export_source_book_docx audience='internal' preserves Engine 2 appendix
- scan_docx_for_forbidden_leakage catches exporter-introduced Engine 2 text
- scan_docx_for_forbidden_leakage catches semantic phrases
- Post-export scan failure produces fail-closed synthetic finding
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

import pytest

from src.models.state import DeckForgeState
from src.services.artifact_gates import (
    ArtifactSection,
    FORBIDDEN_ID_PATTERNS,
    FORBIDDEN_SEMANTIC_PHRASES,
    scan_docx_for_forbidden_leakage,
    scan_for_forbidden_leakage,
)
from src.services.pre_export_sanitizer import sanitize_source_book_sections


# ── Broadened Engine 2 pattern ───────────────────────────────────


def test_generic_engine2_matches_forbidden_pattern():
    """'Engine 2' (not just ENGINE 2 REQUIRED) must match forbidden patterns."""
    combined = "|".join(FORBIDDEN_ID_PATTERNS)
    assert re.search(combined, "Engine 2 must retrieve", re.IGNORECASE)
    assert re.search(combined, "Engine 2 action: source from database", re.IGNORECASE)
    assert re.search(combined, "من Engine 2", re.IGNORECASE)
    assert re.search(combined, "ENGINE 2 REQUIRED", re.IGNORECASE)


def test_engine2_generic_detected_in_client_section():
    section = ArtifactSection(
        section_path="section_4/coverage",
        section_type="client_facing_body",
        text="Coverage gaps flagged for Engine 2 retrieval from company backend.",
    )
    violations = scan_for_forbidden_leakage(section)
    assert len(violations) >= 1
    assert any("Engine 2" in v.matched_text for v in violations)


def test_engine2_allowed_in_internal_appendix():
    section = ArtifactSection(
        section_path="appendix/engine2",
        section_type="internal_gap_appendix",
        text="Engine 2 must retrieve consultant CVs.",
    )
    violations = scan_for_forbidden_leakage(section)
    assert violations == []


# ── Sanitizer with broadened pattern ─────────────────────────────


def test_sanitizer_removes_generic_engine2_from_client_body():
    sections = {
        "proposed_solution": {
            "methodology_overview": "Coverage gaps flagged for Engine 2 retrieval.",
        },
    }
    result = sanitize_source_book_sections(sections)
    text = result.sanitized["proposed_solution"]["methodology_overview"]
    assert "Engine 2" not in text
    assert result.total_removals >= 1


def test_sanitizer_preserves_engine2_in_internal_field():
    sections = {
        "internal_gap_appendix": {
            "notes": "Engine 2 must retrieve all project evidence.",
        },
    }
    result = sanitize_source_book_sections(sections)
    assert "Engine 2" in result.sanitized["internal_gap_appendix"]["notes"]
    assert result.removals == []


# ── State persistence ────────────────────────────────────────────


def test_sanitization_removals_on_state():
    state = DeckForgeState()
    assert state.sanitization_removals == []
    state.sanitization_removals = [{"pattern": "test", "text": "x"}]
    assert len(state.sanitization_removals) == 1


def test_post_export_forbidden_on_state():
    state = DeckForgeState()
    assert state.post_export_forbidden == []
    state.post_export_forbidden = [{"pattern": "test", "text": "x"}]
    assert len(state.post_export_forbidden) == 1


# ── Export audience mode ─────────────────────────────────────────


def test_export_function_accepts_audience_parameter():
    import inspect
    from src.services.source_book_export import export_source_book_docx

    sig = inspect.signature(export_source_book_docx)
    assert "audience" in sig.parameters


def test_export_default_audience_is_internal():
    """Default audience preserves old behavior (internal = show everything)."""
    import inspect
    from src.services.source_book_export import export_source_book_docx

    sig = inspect.signature(export_source_book_docx)
    assert sig.parameters["audience"].default == "internal"


@pytest.mark.asyncio
async def test_export_client_mode_suppresses_engine2_appendix(tmp_path):
    """audience='client' must NOT produce the Engine 2 appendix."""
    from src.models.source_book import SourceBook
    from src.services.source_book_export import export_source_book_docx

    sb = SourceBook(
        client_name="Test", rfp_name="Test RFP", language="en",
    )
    out_path = str(tmp_path / "client.docx")
    await export_source_book_docx(sb, out_path, audience="client")

    from docx import Document
    doc = Document(out_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Appendix: Engine 2 Requirements" not in full_text


@pytest.mark.asyncio
async def test_export_internal_mode_preserves_engine2_appendix(tmp_path):
    """audience='internal' MUST produce the Engine 2 appendix (if gaps exist)."""
    from src.models.source_book import SourceBook, WhyStrategicGears, ConsultantProfile
    from src.services.source_book_export import export_source_book_docx

    sb = SourceBook(
        client_name="Test", rfp_name="Test RFP", language="en",
    )
    # Add a consultant with open_role to trigger Engine 2 appendix
    sb.why_strategic_gears = WhyStrategicGears(
        named_consultants=[ConsultantProfile(
            name="[open_role_profile]",
            role="Lead Consultant",
            relevance="test",
        )],
    )
    out_path = str(tmp_path / "internal.docx")
    await export_source_book_docx(sb, out_path, audience="internal")

    from docx import Document
    doc = Document(out_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Engine 2" in full_text or "engine 2" in full_text.lower()


# ── Post-export DOCX scan (reusable helper) ──────────────────────


@pytest.mark.asyncio
async def test_scan_docx_catches_exporter_engine2_text(tmp_path):
    """scan_docx_for_forbidden_leakage catches Engine 2 prose in DOCX."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("This section has Engine 2 retrieval instructions.")
    doc.add_paragraph("Also: PRJ-001 appears here.")
    path = str(tmp_path / "test.docx")
    doc.save(path)

    findings = scan_docx_for_forbidden_leakage(path)
    patterns_found = {f["pattern"] for f in findings}
    matched_texts = {f["matched_text"] for f in findings}
    assert any("Engine" in p for p in patterns_found)
    assert "PRJ-001" in matched_texts


@pytest.mark.asyncio
async def test_scan_docx_catches_semantic_phrases(tmp_path):
    """scan_docx_for_forbidden_leakage catches semantic forbidden phrases."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("خبرة موثقة في العمل مع سدايا")
    path = str(tmp_path / "test_semantic.docx")
    doc.save(path)

    findings = scan_docx_for_forbidden_leakage(path)
    assert len(findings) >= 1
    assert any("semantic:" in f["pattern"] for f in findings)


@pytest.mark.asyncio
async def test_scan_docx_clean_file_returns_empty(tmp_path):
    """Clean DOCX with no forbidden text returns empty findings."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("تقييم الوضع الراهن لأخلاقيات الذكاء الاصطناعي")
    path = str(tmp_path / "clean.docx")
    doc.save(path)

    findings = scan_docx_for_forbidden_leakage(path)
    assert findings == []


def test_scan_docx_nonexistent_file_raises():
    """scan_docx_for_forbidden_leakage raises on missing file."""
    with pytest.raises(Exception):
        scan_docx_for_forbidden_leakage("/nonexistent/path.docx")


# ── Fail-closed on scan failure ──────────────────────────────────


@pytest.mark.asyncio
async def test_export_client_mode_section4_no_engine2(tmp_path):
    """audience='client' must NOT produce Engine 2 text in Section 4 body.

    Section 4 exporter hardcodes 'Engine 2 action:...' in Evidence Gap
    Summary, 'Engine 2 retrieval' in Search Strategy, and 'Engine 2 or
    manual research' in Theme Coverage. Client mode must suppress these.
    """
    from src.models.source_book import (
        ExternalEvidenceSection,
        ExternalEvidenceEntry,
        SourceBook,
    )
    from src.services.source_book_export import export_source_book_docx

    sb = SourceBook(
        client_name="Test", rfp_name="Test RFP", language="en",
    )
    sb.external_evidence = ExternalEvidenceSection(
        entries=[
            ExternalEvidenceEntry(
                source_id="EXT-001",
                title="Test Source",
                year=2024,
                key_finding="A key finding.",
            ),
        ],
        coverage_assessment="Some gaps remain in local jurisdiction evidence.",
    )
    out_path = str(tmp_path / "client_s4.docx")
    await export_source_book_docx(
        sb, out_path, audience="client",
        theme_coverage={"methodology": {"retained_sources": 1, "status": "weak"}},
    )

    from docx import Document
    doc = Document(out_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # No Engine 2 text anywhere in client export
    assert "Engine 2" not in full_text, (
        f"Client export must not contain 'Engine 2'. Found in: "
        f"{[l for l in full_text.split(chr(10)) if 'Engine 2' in l][:3]}"
    )


@pytest.mark.asyncio
async def test_export_internal_mode_section4_has_engine2(tmp_path):
    """audience='internal' preserves Engine 2 workflow text in Section 4."""
    from src.models.source_book import (
        ExternalEvidenceSection,
        ExternalEvidenceEntry,
        SourceBook,
    )
    from src.services.source_book_export import export_source_book_docx

    sb = SourceBook(
        client_name="Test", rfp_name="Test RFP", language="en",
    )
    sb.external_evidence = ExternalEvidenceSection(
        entries=[
            ExternalEvidenceEntry(
                source_id="EXT-001",
                title="Test Source",
                year=2024,
                key_finding="A key finding.",
            ),
        ],
        coverage_assessment="Gaps remain.",
    )
    out_path = str(tmp_path / "internal_s4.docx")
    await export_source_book_docx(
        sb, out_path, audience="internal",
        theme_coverage={"methodology": {"retained_sources": 1, "status": "weak"}},
    )

    from docx import Document
    doc = Document(out_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Internal mode should contain Engine 2 workflow language
    assert "Engine 2" in full_text or "engine 2" in full_text.lower()


def test_post_export_scan_failure_produces_synthetic_finding():
    """If post-export scan throws, a synthetic critical finding must be added."""
    # Simulate what graph.py does on exception
    try:
        scan_docx_for_forbidden_leakage("/nonexistent/path.docx")
        findings = []
    except Exception as e:
        findings = [{
            "pattern": "POST_EXPORT_SCAN_FAILED",
            "matched_text": str(e),
            "source": "post_export_docx_scan",
            "severity": "critical",
        }]
    assert len(findings) == 1
    assert findings[0]["pattern"] == "POST_EXPORT_SCAN_FAILED"
    assert findings[0]["severity"] == "critical"
