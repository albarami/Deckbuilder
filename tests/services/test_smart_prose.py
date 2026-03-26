"""Tests for DOCX smart-prose rendering (_add_smart_prose, _parse_pipe_table).

Verifies:
1. Plain prose renders as paragraphs
2. Pipe-delimited tables render as real Word tables
3. Mixed content (prose + table + prose) splits correctly
4. Invalid/incomplete tables fall back to prose
5. Separator rows are stripped from parsed tables
6. Header row is bold in rendered Word tables
"""

from __future__ import annotations

from docx import Document

from src.services.source_book_export import (
    _add_smart_prose,
    _add_word_table,
    _parse_pipe_table,
)


# ── _parse_pipe_table ─────────────────────────────────────────────


class TestParsePipeTable:
    """Parse pipe-delimited text lines into row data."""

    def test_basic_table(self):
        lines = [
            "| Header A | Header B |",
            "|----------|----------|",
            "| Cell 1   | Cell 2   |",
            "| Cell 3   | Cell 4   |",
        ]
        rows = _parse_pipe_table(lines)
        assert rows is not None
        assert len(rows) == 3  # header + 2 data rows (separator stripped)
        assert rows[0] == ["Header A", "Header B"]
        assert rows[1] == ["Cell 1", "Cell 2"]

    def test_separator_stripped(self):
        lines = [
            "| A | B |",
            "| --- | --- |",
            "| 1 | 2 |",
        ]
        rows = _parse_pipe_table(lines)
        assert rows is not None
        # Separator should be gone
        assert len(rows) == 2
        assert rows[0] == ["A", "B"]
        assert rows[1] == ["1", "2"]

    def test_single_row_returns_none(self):
        lines = ["| Only Row |"]
        assert _parse_pipe_table(lines) is None

    def test_non_pipe_lines_ignored(self):
        lines = [
            "| A | B |",
            "Just text",
            "| 1 | 2 |",
        ]
        rows = _parse_pipe_table(lines)
        assert rows is not None
        assert len(rows) == 2

    def test_empty_input(self):
        assert _parse_pipe_table([]) is None

    def test_three_column_table(self):
        lines = [
            "| ID | Requirement | Status |",
            "|---|---|---|",
            "| COMP-001 | Security | Pass |",
            "| COMP-002 | Compliance | Pass |",
        ]
        rows = _parse_pipe_table(lines)
        assert rows is not None
        assert len(rows) == 3
        assert len(rows[0]) == 3
        assert rows[1][0] == "COMP-001"


# ── _add_word_table ───────────────────────────────────────────────


class TestAddWordTable:
    """Real Word table creation from row data."""

    def test_creates_table_with_correct_dimensions(self):
        doc = Document()
        rows = [["A", "B"], ["1", "2"], ["3", "4"]]
        _add_word_table(doc, rows)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3
        assert len(table.columns) == 2

    def test_header_row_is_bold(self):
        doc = Document()
        rows = [["Header"], ["Data"]]
        _add_word_table(doc, rows)
        header_cell = doc.tables[0].rows[0].cells[0]
        for para in header_cell.paragraphs:
            for run in para.runs:
                assert run.bold is True

    def test_empty_rows_noop(self):
        doc = Document()
        _add_word_table(doc, [])
        assert len(doc.tables) == 0


# ── _add_smart_prose ──────────────────────────────────────────────


class TestAddSmartProse:
    """Smart prose renderer: plain text → paragraphs, pipes → tables."""

    def test_plain_text_becomes_paragraph(self):
        doc = Document()
        _add_smart_prose(doc, "Hello world.\nSecond line.")
        # Should have at least one paragraph with the text
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Hello world" in t for t in texts)

    def test_pipe_table_becomes_real_table(self):
        doc = Document()
        text = (
            "| Req | Status |\n"
            "|---|---|\n"
            "| COMP-001 | Pass |\n"
            "| COMP-002 | Fail |"
        )
        _add_smart_prose(doc, text)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 3  # header + 2 data

    def test_mixed_prose_and_table(self):
        doc = Document()
        text = (
            "Introduction paragraph.\n"
            "\n"
            "| Col A | Col B |\n"
            "|---|---|\n"
            "| X | Y |\n"
            "\n"
            "Conclusion paragraph."
        )
        _add_smart_prose(doc, text)
        assert len(doc.tables) == 1
        # Should have paragraphs before and after the table
        prose_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        assert "Introduction paragraph." in prose_texts
        assert "Conclusion paragraph." in prose_texts

    def test_empty_text_noop(self):
        doc = Document()
        _add_smart_prose(doc, "")
        assert len(doc.tables) == 0
        assert all(p.text.strip() == "" for p in doc.paragraphs)

    def test_invalid_table_falls_back_to_prose(self):
        doc = Document()
        text = "| Only one row |"
        _add_smart_prose(doc, text)
        # Not enough rows for a table, should be prose
        assert len(doc.tables) == 0

    def test_multiple_tables_in_one_block(self):
        doc = Document()
        text = (
            "| A | B |\n"
            "|---|---|\n"
            "| 1 | 2 |\n"
            "\n"
            "Middle text.\n"
            "\n"
            "| C | D |\n"
            "|---|---|\n"
            "| 3 | 4 |"
        )
        _add_smart_prose(doc, text)
        assert len(doc.tables) == 2

    def test_compliance_table_realistic(self):
        doc = Document()
        text = (
            "Compliance mapping:\n"
            "\n"
            "| Requirement | RFP Clause | SG Response | Evidence | Status |\n"
            "|---|---|---|---|---|\n"
            "| Security | 4.1 | ISO 27001 certified | CLM-001 | Compliant |\n"
            "| Data Protection | 4.2 | PDPL compliant | CLM-002 | Compliant |\n"
            "| Team Quals | 5.1 | PMP certified PM | CLM-003 | Compliant |"
        )
        _add_smart_prose(doc, text)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 4  # header + 3 data
        assert len(table.columns) == 5
