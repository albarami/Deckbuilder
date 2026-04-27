"""Regression fixture loader for frozen failed sessions."""
from __future__ import annotations

import json
from pathlib import Path

_FIXTURE_DIR = Path(__file__).parent


def load_conformance_report(session_id: str) -> dict:
    return json.loads(
        (_FIXTURE_DIR / session_id / "conformance_report.json").read_text(
            encoding="utf-8"
        )
    )


def load_evidence_ledger(session_id: str) -> dict:
    return json.loads(
        (_FIXTURE_DIR / session_id / "evidence_ledger.json").read_text(
            encoding="utf-8"
        )
    )


def load_slide_blueprints(session_id: str) -> list:
    return json.loads(
        (_FIXTURE_DIR / session_id / "slide_blueprint_from_source_book.json").read_text(
            encoding="utf-8"
        )
    )


def load_docx_text(session_id: str) -> str:
    from docx import Document

    d = Document(str(_FIXTURE_DIR / session_id / "source_book.docx"))
    parts: list[str] = []
    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)
